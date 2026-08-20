# main.py
import os
import json
import logging
import hashlib
import secrets
import numpy as np
import webbrowser
from datetime import datetime, date, timedelta
from collections import defaultdict
from flask import Flask, render_template, request, redirect, url_for, flash, session, jsonify, send_from_directory
from werkzeug.middleware.proxy_fix import ProxyFix
from functools import wraps
from werkzeug.utils import secure_filename
import joblib
from database import db
import pytesseract
from pdfminer.high_level import extract_text as pdf_extract_text
from PIL import Image

# Configure logging to see important startup messages
logging.basicConfig(level=logging.INFO, format='%(levelname)s:%(message)s')

# Create the app
app = Flask(__name__)
# IMPORTANT: Ensure this secret key is set for sessions and flash messages to work
app.secret_key = os.environ.get("SESSION_SECRET", "dev-secret-key-change-in-production")
app.wsgi_app = ProxyFix(app.wsgi_app, x_proto=1, x_host=1)

# Configure the database
database_url = os.environ.get("DATABASE_URL")
if not database_url or database_url.strip() == "":
    # Use SQLite file named ivf_tracker.db in the project folder
    database_url = "sqlite:///ivf_tracker.db"
    
app.config["SQLALCHEMY_DATABASE_URI"] = database_url
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["UPLOAD_FOLDER"] = "uploads"
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16MB max file size

# Initialize the app with the extension
db.init_app(app)

# Create upload directory if it doesn't exist
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)
os.makedirs("models", exist_ok=True)

# --- Import Models ---
from models import (
    User, PatientData, IVFCycle, WellnessLog, MedicationReminder, ChatMessage,
    MedicalDocument, Prediction, CycleNote, Clinic, MedicalActivity, Message,
    AuditLog, Notification, DocumentVersion, DocumentAnnotation, DocumentShare,
    SimulationResult, PartnerConnection, PartnerInvitation, PartnerSharingPermission,
    SharedTask, SharedNote, PartnerCheckIn, PartnerNotification, PartnerMessage
)

# --- Import AI and Prediction Services ---
from openai_service import (
    get_chatbot_response,
    get_chatbot_response_gemini,
    generate_tts_audio,
    generate_medical_image,
    generate_diagram,
    get_nutrition_plan,
    get_nutrition_analysis,
    get_yoga_routine,
)
from prediction_service import (
    calculate_ivf_success_prediction,
    calculate_embryo_quality_score,
    generate_personalized_protocol
)

# --- Load ML Models ---
try:
    ivf_model = joblib.load("models/ivf_success_model.pkl")
    mood_model = joblib.load("models/mood_trend_model.pkl")
    emotion_model = joblib.load("models/emotion_model.pkl")
    vectorizer = joblib.load("models/emotion_vectorizer.pkl")
    logging.info("All ML models loaded successfully.")
except FileNotFoundError as e:
    logging.warning(f"ML model file not found: {e}. API endpoints will not work.")
    logging.warning("Please run 'python train_models.py' to train and create the models.")
    ivf_model = None
    mood_model = None
    emotion_model = None
    vectorizer = None
except Exception as e:
    logging.error(f"An error occurred while loading ML models: {e}")

# --- Define Additional Models (if not in models.py) ---
class ForumPost(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(100), nullable=False)
    content = db.Column(db.Text, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    user = db.relationship('User', backref=db.backref('posts', lazy=True))

with app.app_context():
    # This command creates the database file and all tables from models.py
    db.create_all()
    logging.info("Database tables created successfully")

    # --- Data Seeding for Demonstration ---
    # This runs once when the app starts if the database is empty.
    if not User.query.first():
        logging.info("Database is empty. Seeding initial data...")

        # Create Clinics FIRST so we can assign the doctor
        # US Clinics
        us_clinic1 = Clinic(name="Hope Fertility Center", city="New York", state="NY")
        us_clinic2 = Clinic(name="Pinnacle IVF", city="Los Angeles", state="CA")
        us_clinic3 = Clinic(name="Aspire Fertility Institute", city="Chicago", state="IL")
        us_clinic4 = Clinic(name="Genesis Fertility Clinic", city="Houston", state="TX")
        
        # Bangalore Clinics
        bng_clinic1 = Clinic(name="Nova IVF Fertility", city="Bangalore", state="KA", description="Comprehensive fertility center offering advanced IVF treatments and personalized care.")
        bng_clinic2 = Clinic(name="Oasis Fertility", city="Bangalore", state="KA", description="Specialized fertility clinic providing cutting-edge reproductive technologies and support.")
        bng_clinic3 = Clinic(name="Iswarya Fertility Centre", city="Bangalore", state="KA", description="Leading fertility center with experienced specialists and state-of-the-art facilities.")
        bng_clinic4 = Clinic(name="Gaudium IVF Centre", city="Bangalore", state="KA", description="Dedicated to helping couples achieve their dream of parenthood through innovative IVF solutions.")
        bng_clinic5 = Clinic(name="Indira IVF", city="Bangalore", state="KA", description="Trusted fertility clinic offering comprehensive reproductive health services.")

        # More Karnataka Clinics
        mys_clinic1 = Clinic(name="Gunasheela Fertility Centre", city="Mysore", state="KA")
        mng_clinic1 = Clinic(name="Manipal Fertility", city="Mangalore", state="KA")
        hub_clinic1 = Clinic(name="Oasis Fertility", city="Hubli", state="KA")
        bel_clinic1 = Clinic(name="KLE Fertility Centre", city="Belgaum", state="KA")
        davan_clinic1 = Clinic(name="Shree IVF", city="Davanagere", state="KA")
        shivam_clinic1 = Clinic(name="Malnad Fertility", city="Shivamogga", state="KA")
        udupi_clinic1 = Clinic(name="Mahe Fertility", city="Udupi", state="KA")
        gulb_clinic1 = Clinic(name="Basaveshwar Hospital IVF", city="Gulbarga", state="KA")


        # Mumbai Clinics
        mum_clinic1 = Clinic(name="Bloom IVF Center", city="Mumbai", state="MH", address="123 Marine Drive", phone="+91-22-1234-5678", website="https://bloomivf.com", description="Leading IVF center in Mumbai with state-of-the-art facilities and experienced fertility specialists.")
        mum_clinic2 = Clinic(name="ART Fertility Clinics", city="Mumbai", state="MH", address="456 Bandra West", phone="+91-22-2345-6789", website="https://artfertility.com", description="Comprehensive fertility treatment center offering advanced ART procedures.")
        mum_clinic3 = Clinic(name="Bavishi Fertility Institute", city="Mumbai", state="MH", address="789 Andheri East", phone="+91-22-3456-7890", website="https://bavishifertility.com", description="Renowned fertility institute with high success rates in IVF treatments.")
        mum_clinic4 = Clinic(name="Progenesis IVF", city="Mumbai", state="MH", address="101 Lower Parel", phone="+91-22-4567-8901", website="https://progenesisivf.com", description="Modern fertility center specializing in personalized IVF protocols.")

        # Delhi Clinics
        delhi_clinic1 = Clinic(name="International Fertility Centre", city="Delhi", state="DL")
        delhi_clinic2 = Clinic(name="Medicover Fertility", city="Delhi", state="DL")
        delhi_clinic3 = Clinic(name="Aakash Healthcare", city="Delhi", state="DL")

        # Chennai Clinics
        chennai_clinic1 = Clinic(name="Prashanth Fertility Research Centre", city="Chennai", state="TN")
        chennai_clinic2 = Clinic(name="Jananam Fertility Centre", city="Chennai", state="TN")

        # Hyderabad Clinics
        hyd_clinic1 = Clinic(name="KIMS Cuddles", city="Hyderabad", state="TS")
        hyd_clinic2 = Clinic(name="Srujana Fertility Centre", city="Hyderabad", state="TS")

        db.session.add_all([us_clinic1, us_clinic2, us_clinic3, us_clinic4, 
                            bng_clinic1, bng_clinic2, bng_clinic3, bng_clinic4, bng_clinic5,
                            mys_clinic1, mng_clinic1, hub_clinic1, bel_clinic1, davan_clinic1, 
                            shivam_clinic1, udupi_clinic1, gulb_clinic1,
                            mum_clinic1, mum_clinic2, mum_clinic3, mum_clinic4,
                            delhi_clinic1, delhi_clinic2, delhi_clinic3,
                            chennai_clinic1, chennai_clinic2, hyd_clinic1, hyd_clinic2])
        db.session.commit() # Commit to get clinic IDs

        # Create Admin User
        admin_user = User(
            username="admin",
            email="admin@example.com",
            first_name="Admin",
            last_name="User",
            user_type="admin"
        )
        admin_user.set_password("admin123")
        db.session.add(admin_user)

        # Create Doctor User
        doctor_user = User(
            username="doctor",
            email="doctor@example.com",
            first_name="John",
            last_name="Doe",
            user_type="doctor",
            clinic_id=us_clinic1.id # Assign doctor to a clinic
        )
        doctor_user.set_password("doctor123")
        db.session.add(doctor_user)
        
        db.session.commit()
        logging.info("Initial data seeded successfully.")

# --- Routes ---

# --- Security Helper Functions ---
def log_action(user_id, action, details=None):
    """Logs sensitive user actions to the database for audit trails."""
    try:
        ip = request.remote_addr
        log = AuditLog(user_id=user_id, action=action, details=details, ip_address=ip)
        db.session.add(log)
        db.session.commit()
    except Exception as e:
        app.logger.error(f"Failed to create audit log: {e}")


def create_notification(user_id, message, type='info', category='system', link=None):
    """Creates a smart notification for a user and logs critical ones."""
    try:
        notification = Notification(
            user_id=user_id,
            message=message,
            type=type,
            category=category,
            link=link
        )
        db.session.add(notification)
        db.session.commit()

        if type == 'critical':
            log_action(user_id, 'NOTIFICATION_CRITICAL', f'Critical alert: {message[:100]}')
        return notification
    except Exception as e:
        app.logger.error(f"Failed to create notification: {e}")
        db.session.rollback()
        return None

# Utility functions
def login_required(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        if 'user_id' not in session:
            flash('Please log in to access this page.', 'error')
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return wrapper

def admin_required(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        if 'user_id' not in session or session.get('user_type') != 'admin':
            flash('You do not have permission to access this page.', 'error')
            return redirect(url_for('index'))
        return f(*args, **kwargs)
    return wrapper

def allowed_file(filename):
    ALLOWED_EXTENSIONS = {'txt', 'pdf', 'png', 'jpg', 'jpeg', 'gif', 'doc', 'docx'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def hash_invitation_token(token):
    return hashlib.sha256(token.encode('utf-8')).hexdigest()


def get_partner_connection_for_user(user):
    if user is None:
        return None
    if user.user_type == 'partner':
        return PartnerConnection.query.filter_by(partner_id=user.id, status='active').first()
    if user.user_type == 'patient':
        return PartnerConnection.query.filter_by(patient_id=user.id, status='active').first()
    return None


def get_partner_permissions(connection):
    if connection is None:
        return None
    permissions = connection.permissions
    if permissions is None:
        permissions = PartnerSharingPermission(connection_id=connection.id)
        db.session.add(permissions)
        db.session.commit()
    return permissions


def find_valid_invitation(token):
    if not token:
        return None
    token_hash = hash_invitation_token(token)
    invitation = PartnerInvitation.query.filter_by(token_hash=token_hash).first()
    if invitation is None:
        return None
    if invitation.status in ['declined', 'accepted', 'cancelled']:
        return None
    if invitation.expires_at < datetime.utcnow():
        invitation.status = 'expired'
        db.session.commit()
        return None
    return invitation


# --- Medical Data Extraction Functions ---
import re

def extract_medical_params(text):
    """Extract IVF-related medical parameters from text using regex patterns"""
    extracted = {
        'age': None,
        'amh': None,
        'fsh': None,
        'lh': None,
        'afc': None,
        'uterus': None,
        'ovaries': None,
        'blood_tests': []
    }
    
    if not text:
        return extracted
    
    text_lower = text.lower()
    
    # Extract Age patterns
    age_patterns = [
        r'age[:\s]+(\d+)',
        r'patient.{0,20}age[:\s]+(\d+)',
        r'aged?\s+(\d+)',
        r'(\d+)\s*years?\s*old'
    ]
    for pattern in age_patterns:
        match = re.search(pattern, text_lower)
        if match:
            extracted['age'] = int(match.group(1))
            break
    
    # Extract AMH patterns
    amh_patterns = [
        r'amh[:\s]+([\d.]+)\s*(ng/ml|ng/mL)?',
        r'anti[-\s]m[ü]llerian[:\s]+([\d.]+)',
        r'amh.*?([\d.]+)'
    ]
    for pattern in amh_patterns:
        match = re.search(pattern, text_lower)
        if match:
            try:
                extracted['amh'] = float(match.group(1))
                break
            except:
                pass
    
    # Extract FSH patterns
    fsh_patterns = [
        r'fsh[:\s]+([\d.]+)\s*(mIU/mL|mlu?/ml)?',
        r'follicle stimulating[:\s]+([\d.]+)',
        r'fsh.*?([\d.]+)'
    ]
    for pattern in fsh_patterns:
        match = re.search(pattern, text_lower)
        if match:
            try:
                extracted['fsh'] = float(match.group(1))
                break
            except:
                pass
    
    # Extract LH patterns
    lh_patterns = [
        r'lh[:\s]+([\d.]+)\s*(mIU/mL|mlu?/ml)?',
        r'luteinizing[:\s]+([\d.]+)',
        r'lh.*?([\d.]+)'
    ]
    for pattern in lh_patterns:
        match = re.search(pattern, text_lower)
        if match:
            try:
                extracted['lh'] = float(match.group(1))
                break
            except:
                pass
    
    # Extract AFC (Antral Follicle Count) patterns
    afc_patterns = [
        r'afc[:\s]+(\d+)',
        r'antral\s*follicle\s*count[:\s]+(\d+)',
        r'total\s*antral\s*follicle[:\s]+(\d+)',
        r'(\d+)\s*antral\s*follicles?'
    ]
    for pattern in afc_patterns:
        match = re.search(pattern, text_lower)
        if match:
            extracted['afc'] = int(match.group(1))
            break
    
    # Extract Uterus condition patterns
    uterus_patterns = [
        r'uterus[:\s]+(normal|abnormal|good|fair|poor)',
        r'uterine[:\s]+(normal|abnormal|good|fair|poor)',
        r'endometrium[:\s]+(normal|abnormal|good|fair|poor)'
    ]
    for pattern in uterus_patterns:
        match = re.search(pattern, text_lower)
        if match:
            extracted['uterus'] = match.group(1).capitalize()
            break
    
    # Extract Ovary condition patterns
    ovary_patterns = [
        r'ovaries[:\s]+(normal|abnormal|good|fair|poor)',
        r'ovarian[:\s]+(normal|abnormal|good|fair|poor)',
        r'right\s*ovary[:\s]+(normal|abnormal)',
        r'left\s*ovary[:\s]+(normal|abnormal)'
    ]
    for pattern in ovary_patterns:
        match = re.search(pattern, text_lower)
        if match:
            extracted['ovaries'] = match.group(1).capitalize()
            break
    
    # Extract blood test results
    blood_params = {
        'tsh': r'tsh[:\s]+([\d.]+)',
        'prolactin': r'prolactin[:\s]+([\d.]+)',
        'estradiol': r'estradiol[:\s]+([\d.]+)',
        'progesterone': r'progesterone[:\s]+([\d.]+)',
        'testosterone': r'testosterone[:\s]+([\d.]+)',
        'hemoglobin': r'hemoglobin[:\s]+([\d.]+)',
        'glucose': r'glucose[:\s]+([\d.]+)'
    }
    
    for param, pattern in blood_params.items():
        match = re.search(pattern, text_lower)
        if match:
            try:
                extracted['blood_tests'].append({
                    'name': param.capitalize(),
                    'value': float(match.group(1))
                })
            except:
                pass
    
    return extracted


def generate_medical_insights(params):
    """Generate automatic insights based on extracted medical parameters"""
    insights = []
    
    if params.get('age') is not None:
        age = params['age']
        if age < 35:
            insights.append("Patient is in the optimal reproductive age group (under 35).")
        elif age < 40:
            insights.append("Patient is in the advanced maternal age group. Special protocols may be needed.")
        else:
            insights.append("Patient is in advanced maternal age. Consider discussing donor eggs or other options.")
    
    if params.get('amh') is not None:
        amh = params['amh']
        if amh >= 2:
            insights.append("AMH level indicates good ovarian reserve.")
        elif amh >= 1:
            insights.append("AMH level is in the low-normal range.")
        else:
            insights.append("AMH level indicates diminished ovarian reserve. Consider aggressive stimulation protocols.")
    
    if params.get('afc') is not None:
        afc = params['afc']
        if afc >= 12:
            insights.append("Good Antral Follicle Count indicates healthy ovarian response.")
        elif afc >= 6:
            insights.append("Moderate Antral Follicle Count.")
        else:
            insights.append("Low Antral Follicle Count. May need customized protocols.")
    
    if params.get('uterus') and params.get('uterus').lower() == 'normal':
        insights.append("Uterine conditions appear normal for IVF transfer.")
    
    if params.get('ovaries') and params.get('ovaries').lower() == 'normal':
        insights.append("Ovarian conditions appear normal.")
    
    if params.get('fsh') is not None:
        fsh = params['fsh']
        if fsh > 10:
            insights.append("Elevated FSH may indicate diminished ovarian reserve.")
        elif fsh < 5:
            insights.append("Low FSH levels. May need further evaluation.")
    
    if not insights:
        insights.append("Continue with standard IVF protocol based on available data.")
    
    return insights[:3]

# Authentication routes
@app.route('/')
def index():
    if 'user_id' in session:
        user = User.query.get(session['user_id'])
        if user.user_type == 'admin':
            return redirect(url_for('admin_dashboard'))
        elif user.user_type == 'doctor':
            return redirect(url_for('doctor_dashboard'))
        elif user.user_type == 'partner':
            return redirect(url_for('partner_dashboard'))
        else:
            return redirect(url_for('patient_dashboard'))
    return render_template('index.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    login_type = request.args.get('type', 'general')  # Get type from query param, default to 'general'

    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']

        user = User.query.filter_by(email=email).first()
        if user and user.check_password(password):
            session['user_id'] = user.id
            session['user_type'] = user.user_type
            flash(f'Welcome back, {user.first_name}!', 'success')
            
            # Security: Log the login action
            log_action(user.id, 'LOGIN', 'User logged in successfully')

            if user.user_type == 'admin':
                return redirect(url_for('admin_dashboard'))
            elif user.user_type == 'doctor':
                return redirect(url_for('doctor_dashboard'))
            elif user.user_type == 'partner':
                return redirect(url_for('partner_dashboard'))
            else:
                return redirect(url_for('patient_dashboard'))
        else:
            flash('Invalid email or password.', 'error')

    return render_template('login.html', login_type=login_type)

@app.route("/reset_password", methods=['GET', 'POST'])
def reset_request():
    if 'user_id' in session:
        return redirect(url_for('index'))
    if request.method == 'POST':
        email = request.form.get('email')
        user = User.query.filter_by(email=email).first()
        if user:
            token = user.get_reset_token()
            reset_link = url_for('reset_token', token=token, _external=True)
            
            # In a real app, you would email this link.
            # For this project, we'll print it to the console.
            print("--- PASSWORD RESET LINK ---")
            print(f"To reset the password for {user.email}, use this link:")
            print(reset_link)
            print("---------------------------")
            
            flash('A password reset link has been generated. Please check the console for the link.', 'info')
        else:
            flash('No account found with that email address.', 'warning')
        return redirect(url_for('login'))
    return render_template('request_reset.html', title='Reset Password')

@app.route("/reset_password/<token>", methods=['GET', 'POST'])
def reset_token(token):
    if 'user_id' in session:
        return redirect(url_for('index'))
    user = User.verify_reset_token(token)
    if user is None:
        flash('That is an invalid or expired token.', 'warning')
        return redirect(url_for('reset_request'))
    if request.method == 'POST':
        password = request.form.get('password')
        user.set_password(password)
        db.session.commit()
        flash('Your password has been updated! You are now able to log in.', 'success')
        return redirect(url_for('login'))
    return render_template('reset_password.html', title='Reset Password')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        # Security: Consent Checkbox Validation
        if not request.form.get('agree_terms'):
            flash('You must agree to the privacy policy and data usage terms to register.', 'error')
            return redirect(url_for('register'))

        try:
            # Check if username or email already exists
            if User.query.filter_by(username=request.form['username']).first() or \
               User.query.filter_by(email=request.form['email']).first():
                flash('Username or Email already registered.', 'error')
                return redirect(url_for('register'))

            user = User(
                username=request.form['username'],
                email=request.form['email'],
                first_name=request.form['first_name'],
                last_name=request.form['last_name'],
                user_type='patient', # Default to patient for this route
                phone=request.form.get('phone'),
                date_of_birth=datetime.strptime(request.form['date_of_birth'], '%Y-%m-%d').date() if request.form.get('date_of_birth') else None
            )
            user.set_password(request.form['password'])
            
            db.session.add(user)
            db.session.commit()
            
            # Create patient data record if user is a patient
            if user.user_type == 'patient':
                patient_data = PatientData(user_id=user.id)
                db.session.add(patient_data)
                db.session.commit()
            
            # Security: Log registration
            log_action(user.id, 'REGISTER', f'New patient registration: {user.email}')
            
            flash('Registration successful! Please log in.', 'success')
            return redirect(url_for('login'))
            
        except Exception as e:
            db.session.rollback()
            # Log the full error for debugging
            app.logger.error(f"Registration failed: {e}")
            flash('Registration failed due to a server error.', 'error')
    
    return render_template('register.html')


@app.route('/register/partner', methods=['GET', 'POST'])
def register_partner():
    token = request.args.get('token') or request.form.get('token')
    invitation = find_valid_invitation(token) if token else None

    if request.method == 'POST':
        if not request.form.get('agree_terms'):
            flash('You must agree to the privacy policy and data usage terms to register.', 'error')
            return redirect(url_for('register_partner', token=token))

        full_name = request.form.get('full_name', '').strip()
        email = request.form.get('email', '').strip().lower()
        password = request.form.get('password', '')
        confirm_password = request.form.get('confirm_password', '')
        relationship = request.form.get('relationship', 'Partner')

        if not full_name or not email or not password:
            flash('Please complete all required fields.', 'error')
            return redirect(url_for('register_partner', token=token))

        if password != confirm_password:
            flash('Passwords do not match.', 'error')
            return redirect(url_for('register_partner', token=token))

        if User.query.filter_by(email=email).first():
            flash('An account with this email already exists. Please log in.', 'error')
            return redirect(url_for('login'))

        names = full_name.split(maxsplit=1)
        first_name = names[0]
        last_name = names[1] if len(names) > 1 else 'Partner'

        partner_user = User(
            username=f"partner_{email.split('@')[0]}_{secrets.token_hex(4)}",
            email=email,
            first_name=first_name,
            last_name=last_name,
            user_type='partner',
            phone=request.form.get('phone'),
            date_of_birth=datetime.strptime(request.form['date_of_birth'], '%Y-%m-%d').date() if request.form.get('date_of_birth') else None,
        )
        partner_user.set_password(password)
        db.session.add(partner_user)
        db.session.commit()

        if invitation is not None:
            existing_connection = PartnerConnection.query.filter_by(patient_id=invitation.patient_id, partner_id=partner_user.id, status='active').first()
            if not existing_connection:
                connection = PartnerConnection(patient_id=invitation.patient_id, partner_id=partner_user.id, status='active', accepted_at=datetime.utcnow())
                db.session.add(connection)
                db.session.commit()
                permissions = PartnerSharingPermission(
                    connection_id=connection.id,
                    treatment_shared=True,
                    appointments_shared=True,
                    medications_shared=True,
                    wellness_shared=False,
                    mood_shared=False,
                    nutrition_shared=True,
                    documents_shared=False,
                    doctor_notes_shared=False,
                    messages_enabled=True,
                    tasks_enabled=True,
                )
                db.session.add(permissions)
            invitation.status = 'accepted'
            invitation.accepted_at = datetime.utcnow()
            db.session.commit()

        session['user_id'] = partner_user.id
        session['user_type'] = 'partner'
        flash('Partner account created successfully. Welcome to the partner dashboard.', 'success')
        return redirect(url_for('partner_dashboard'))

    return render_template('partner_register.html', token=token, invitation=invitation)

@app.route('/register/doctor', methods=['GET', 'POST'])
def register_doctor():
    if request.method == 'POST':
        # Security: Consent Checkbox Validation
        if not request.form.get('agree_terms'):
            flash('You must agree to the privacy policy and data usage terms to register.', 'error')
            return redirect(url_for('register_doctor'))

        try:
            # Check if username or email already exists
            if User.query.filter_by(username=request.form['username']).first() or \
               User.query.filter_by(email=request.form['email']).first():
                flash('Username or Email already registered.', 'error')
                return redirect(url_for('register_doctor'))

            clinic_id = request.form.get('clinic_id')
            if not clinic_id:
                flash('Please select a clinic.', 'error')
                return redirect(url_for('register_doctor'))

            user = User(
                username=request.form['username'],
                email=request.form['email'],
                first_name=request.form['first_name'],
                last_name=request.form['last_name'],
                user_type='doctor',
                phone=request.form.get('phone'),
                date_of_birth=datetime.strptime(request.form.get('date_of_birth'), '%Y-%m-%d').date() if request.form.get('date_of_birth') else None,
                clinic_id=int(clinic_id)
            )
            user.set_password(request.form['password'])
            
            db.session.add(user)
            db.session.commit()
            
            # Security: Log registration
            log_action(user.id, 'REGISTER_DOCTOR', f'New doctor registration: {user.email}')
            
            flash('Doctor registration successful! Please log in.', 'success')
            return redirect(url_for('login', type='doctor'))
            
        except Exception as e:
            db.session.rollback()
            app.logger.error(f"Doctor registration failed: {e}")
            flash('Registration failed due to a server error.', 'error')

    # Fetch clinics for the dropdown
    clinics_query = Clinic.query.order_by(Clinic.state, Clinic.city, Clinic.name).all()
    
    # Group clinics by state for a more organized dropdown
    grouped_clinics = defaultdict(list)
    for clinic in clinics_query:
        grouped_clinics[clinic.state].append(clinic)
        
    # Sort the dictionary by state (key) to ensure alphabetical order in the dropdown
    sorted_grouped_clinics = dict(sorted(grouped_clinics.items()))
        
    return render_template('register_doctor.html', grouped_clinics=sorted_grouped_clinics)

@app.route('/logout')
@login_required
def logout():
    # Security: Log logout before clearing session
    log_action(session['user_id'], 'LOGOUT', 'User logged out')
    session.clear()
    flash('You have been logged out successfully.', 'info')
    return redirect(url_for('index'))

# Dashboard routes
@app.route('/patient_dashboard')
@login_required
def patient_dashboard():
    user = User.query.get(session['user_id'])
    
    # Add a check to handle invalid sessions (e.g., after a database reset)
    if user is None:
        session.clear() # Clear the invalid session
        flash('Your session has expired. Please log in again.', 'error')
        return redirect(url_for('login'))

    # Explicitly check user type and redirect if incorrect
    if user.user_type != 'patient':
        flash('Access denied.', 'error')
        # Redirect to the correct dashboard based on their actual user type
        if user.user_type == 'doctor':
            return redirect(url_for('doctor_dashboard'))
        elif user.user_type == 'admin':
            return redirect(url_for('admin_dashboard'))
        return redirect(url_for('logout')) # Fallback to logout
    
    patient_data = PatientData.query.filter_by(user_id=user.id).first()
    
    # Get latest cycle for progress tracking
    latest_cycle = IVFCycle.query.filter_by(patient_id=user.id).order_by(IVFCycle.start_date.desc()).first()
    
    # Get the latest prediction for the stat card
    latest_prediction = Prediction.query.filter_by(user_id=user.id).order_by(Prediction.prediction_date.desc()).first()
    
    # Get active medications
    active_medications = MedicationReminder.query.filter_by(user_id=user.id, is_active=True).order_by(MedicationReminder.time_of_day).all()
    
    # Get recent wellness logs for the stat card and list
    recent_wellness = WellnessLog.query.filter_by(user_id=user.id).order_by(WellnessLog.date.desc()).limit(7).all()
    # Check if today's wellness log exists
    today_wellness_log = WellnessLog.query.filter_by(user_id=user.id, date=date.today()).first()
    
    # Get a personalized AI tip (mocked for now)
    ai_tip = get_chatbot_response("Give me one personalized wellness tip for today.", user, patient_data)

    partner_connection = PartnerConnection.query.filter_by(patient_id=user.id, status='active').first()
    partner_user = User.query.get(partner_connection.partner_id) if partner_connection else None
    partner_pending = PartnerInvitation.query.filter_by(patient_id=user.id, status='pending').first()
    
    return render_template('patient_dashboard.html', 
                          user=user, 
                          patient_data=patient_data,
                          latest_cycle=latest_cycle,
                          latest_prediction=latest_prediction,
                          recent_wellness=recent_wellness,
                          active_medications=active_medications,
                          today_wellness_log_exists=(today_wellness_log is not None),
                          ai_tip=ai_tip,
                          partner_connection=partner_connection,
                          partner_user=partner_user,
                          partner_pending=partner_pending)

@app.route('/doctor_dashboard')
@login_required
def doctor_dashboard():
    user = User.query.get(session['user_id'])

    # Add a check to handle invalid sessions
    if user is None:
        session.clear()
        flash('Your session has expired. Please log in again.', 'error')
        return redirect(url_for('login'))

    if user.user_type != 'doctor':
        flash('Access denied. This dashboard is for doctors only.', 'error')
        return redirect(url_for('patient_dashboard')) # Redirect non-doctors to the patient dashboard

    # Get all patients
    patients = User.query.filter_by(user_type='patient').all()

    # Get recent cycles
    recent_cycles = IVFCycle.query.order_by(IVFCycle.created_at.desc()).limit(10).all()

    # Get recent medical activities
    recent_activities = MedicalActivity.query.join(User, MedicalActivity.patient_id == User.id).order_by(MedicalActivity.performed_date.desc()).limit(10).all()

    return render_template('doctor_dashboard.html',
                          user=user,
                          patients=patients,
                          recent_cycles=recent_cycles,
                          recent_activities=recent_activities)

@app.route('/patient/<int:patient_id>')
@login_required
def patient_details(patient_id):
    user = User.query.get(session['user_id'])
    if user.user_type != 'doctor':
        flash('Access denied. This page is for doctors only.', 'error')
        return redirect(url_for('patient_dashboard'))

    patient = User.query.get_or_404(patient_id)
    if patient.user_type != 'patient':
        flash('Invalid patient.', 'error')
        return redirect(url_for('doctor_dashboard'))
    
    # Security: Audit Log for accessing patient data (HIPAA requirement)
    log_action(user.id, 'VIEW_PATIENT', f'Doctor viewed patient ID {patient.id}')

    patient_data = PatientData.query.filter_by(user_id=patient.id).first()

    # Get patient's cycles
    cycles = IVFCycle.query.filter_by(patient_id=patient.id).order_by(IVFCycle.start_date.desc()).all()

    # Get recent wellness logs
    recent_wellness = WellnessLog.query.filter_by(user_id=patient.id).order_by(WellnessLog.date.desc()).limit(7).all()

    # Get latest prediction
    latest_prediction = Prediction.query.filter_by(user_id=patient.id).order_by(Prediction.prediction_date.desc()).first()

    # Get active medications
    active_medications = MedicationReminder.query.filter_by(user_id=patient.id, is_active=True).order_by(MedicationReminder.time_of_day).all()

    return render_template('patient_details.html',
                          user=user,
                          patient=patient,
                          patient_data=patient_data,
                          cycles=cycles,
                          recent_wellness=recent_wellness,
                          latest_prediction=latest_prediction,
                          active_medications=active_medications)

@app.route('/edit_patient/<int:patient_id>', methods=['GET', 'POST'])
@login_required
def edit_patient(patient_id):
    user = User.query.get(session['user_id'])
    if user.user_type != 'doctor':
        flash('Access denied. This page is for doctors only.', 'error')
        return redirect(url_for('patient_dashboard'))

    patient = User.query.get_or_404(patient_id)
    if patient.user_type != 'patient':
        flash('Invalid patient.', 'error')
        return redirect(url_for('doctor_dashboard'))

    patient_data = PatientData.query.filter_by(user_id=patient.id).first()

    if not patient_data:
        patient_data = PatientData(user_id=patient.id)
        db.session.add(patient_data)

    if request.method == 'POST':
        # Update user data
        patient.first_name = request.form['first_name']
        patient.last_name = request.form['last_name']
        patient.phone = request.form.get('phone')

        if request.form.get('date_of_birth'):
            patient.date_of_birth = datetime.strptime(request.form['date_of_birth'], '%Y-%m-%d').date()

        # Update patient data with type casting and validation
        if request.form.get('age'):
            patient_data.age = int(request.form['age'])

        # Safely handle float conversions
        height_cm = float(request.form.get('height', 0))
        weight_kg = float(request.form.get('weight', 0))

        if height_cm > 0:
            patient_data.height = height_cm
        if weight_kg > 0:
            patient_data.weight = weight_kg

        # Calculate BMI
        if patient_data.height and patient_data.height > 0 and patient_data.weight:
            patient_data.bmi = patient_data.weight / ((patient_data.height / 100) ** 2)
        else:
            patient_data.bmi = None

        if request.form.get('amh_level'):
            patient_data.amh_level = float(request.form['amh_level'])
        if request.form.get('fsh_level'):
            patient_data.fsh_level = float(request.form['fsh_level'])

        patient_data.diagnosis = request.form.get('diagnosis')
        patient_data.medical_history = request.form.get('medical_history')
        patient_data.medications = request.form.get('medications')
        patient_data.allergies = request.form.get('allergies')
        patient_data.lifestyle_factors = request.form.get('lifestyle_factors')

        if request.form.get('previous_pregnancies'):
            patient_data.previous_pregnancies = int(request.form['previous_pregnancies'])
        if request.form.get('previous_ivf_cycles'):
            patient_data.previous_ivf_cycles = int(request.form['previous_ivf_cycles'])
        if request.form.get('partner_age'):
            patient_data.partner_age = int(request.form['partner_age'])

        patient_data.partner_diagnosis = request.form.get('partner_diagnosis')

        db.session.commit()
        flash('Patient information updated successfully!', 'success')
        return redirect(url_for('patient_details', patient_id=patient.id))

    return render_template('edit_patient.html', user=user, patient=patient, patient_data=patient_data)

# Doctor's personal notes page
@app.route('/my_notes')
@login_required
def my_notes():
    user = User.query.get(session['user_id'])
    if user.user_type != 'doctor':
        flash('Access denied. This page is for doctors only.', 'error')
        return redirect(url_for('patient_dashboard'))

    # Query for all notes written by the current doctor, ordered by most recent
    # Joined with IVFCycle and User to eager load patient info for efficiency
    notes = CycleNote.query.filter_by(doctor_id=user.id).join(IVFCycle).join(User, IVFCycle.patient_id == User.id).order_by(CycleNote.created_at.desc()).all()

    return render_template('my_notes.html', user=user, notes=notes)

# Patient's view of doctor's notes
@app.route('/my_cycle_notes')
@login_required
def my_cycle_notes():
    user = User.query.get(session['user_id'])

    # Add a check to handle invalid sessions (e.g., after a database reset)
    if user is None:
        session.clear() # Clear the invalid session
        flash('Your session has expired. Please log in again.', 'error')
        return redirect(url_for('login'))

    if user.user_type != 'patient':
        flash('Access denied. This page is for patients only.', 'error')
        return redirect(url_for('index'))

    # Query for all notes related to the patient's cycles
    # Eager load related cycle and doctor info for efficiency
    from sqlalchemy.orm import joinedload
    notes = CycleNote.query.join(IVFCycle).filter(IVFCycle.patient_id == user.id).options(
        joinedload(CycleNote.cycle),
        joinedload(CycleNote.doctor)
    ).order_by(CycleNote.created_at.desc()).all()

    return render_template('view_notes.html', user=user, notes=notes)

# Find a Clinic/Doctor page
@app.route('/find_clinic', methods=['GET'])
@login_required
def find_clinic():
    # Get query parameters
    query = request.args.get('query', '').strip()
    state_filter = request.args.get('state', '').strip()
    clinic_type_filter = request.args.get('clinic_type', '').strip()
    page = request.args.get('page', 1, type=int)
    per_page = 20  # Clinics per page

    # Base query
    base_query = Clinic.query

    # Apply filters
    if query:
        # Improved search: fuzzy matching on name, city, state, zip_code
        search_conditions = []
        search_terms = query.split()
        for term in search_terms:
            term_lower = term.lower()
            search_conditions.append(db.func.lower(Clinic.name).like(f"%{term_lower}%"))
            search_conditions.append(db.func.lower(Clinic.city).like(f"%{term_lower}%"))
            search_conditions.append(db.func.lower(Clinic.state).like(f"%{term_lower}%"))
            if Clinic.zip_code is not None:
                search_conditions.append(db.func.lower(Clinic.zip_code).like(f"%{term_lower}%"))
        base_query = base_query.filter(db.or_(*search_conditions))

    if state_filter:
        base_query = base_query.filter(db.func.lower(Clinic.state) == state_filter.lower())

    if clinic_type_filter:
        base_query = base_query.filter(db.func.lower(Clinic.clinic_type) == clinic_type_filter.lower())

    # Order by state, city, name
    base_query = base_query.order_by(Clinic.state, Clinic.city, Clinic.name)

    # Paginate
    clinics_pagination = base_query.paginate(page=page, per_page=per_page, error_out=False)
    clinics = clinics_pagination.items

    # Get filter options
    all_states = db.session.query(Clinic.state).distinct().order_by(Clinic.state).all()
    states = [state[0] for state in all_states]

    all_clinic_types = db.session.query(Clinic.clinic_type).distinct().filter(Clinic.clinic_type.isnot(None)).order_by(Clinic.clinic_type).all()
    clinic_types = [ct[0] for ct in all_clinic_types]

    total_clinics = base_query.count()

    # Convert clinics to dictionaries for JSON serialization
    clinics_data = []
    for clinic in clinics:
        clinics_data.append({
            'id': clinic.id,
            'name': clinic.name,
            'address': clinic.address,
            'city': clinic.city,
            'state': clinic.state,
            'zip_code': clinic.zip_code,
            'phone': clinic.phone,
            'website': clinic.website,
            'latitude': clinic.latitude,
            'longitude': clinic.longitude,
            'description': clinic.description,
            'clinic_type': clinic.clinic_type
        })

    return render_template('find_clinic.html',
                          clinics=clinics,
                          clinics_data=clinics_data,
                          query=query,
                          state_filter=state_filter,
                          clinic_type_filter=clinic_type_filter,
                          clinics_pagination=clinics_pagination,
                          states=states,
                          clinic_types=clinic_types,
                          total_clinics=total_clinics)

# Clinic Detail page
@app.route('/clinic/<int:clinic_id>')
@login_required
def clinic_detail(clinic_id):
    clinic = Clinic.query.get_or_404(clinic_id)
    return render_template('clinic_detail.html', clinic=clinic)

# --- Admin Routes ---
@app.route('/admin')
@admin_required
def admin_dashboard():
    user = User.query.get(session['user_id'])
    user_count = User.query.count()
    patient_count = User.query.filter_by(user_type='patient').count()
    doctor_count = User.query.filter_by(user_type='doctor').count()
    clinic_count = Clinic.query.count()
    
    # --- AI Model Statistics ---
    model_stats = {
        'accuracy': 'N/A',
        'feature_importance': {},
        'labels': [],
        'data': []
    }

    if ivf_model:
        # 1. Feature Importance
        if hasattr(ivf_model, 'feature_importances_'):
            feature_names = ['Age', 'BMI', 'AMH', 'FSH', 'Prev IVF', 'Stress', 'Sleep', 'Exercise']
            importances = ivf_model.feature_importances_
            # Sort by importance
            indices = np.argsort(importances)[::-1]
            model_stats['labels'] = [feature_names[i] for i in indices]
            model_stats['data'] = [round(importances[i], 3) for i in indices]
        
        # 2. Calculate Accuracy on Real Data (Completed Cycles)
        try:
            completed_cycles = IVFCycle.query.filter(
                IVFCycle.outcome.in_(['BFP', 'Live Birth', 'BFN', 'Miscarriage'])
            ).all()
            
            if completed_cycles:
                correct = 0
                total = 0
                
                for cycle in completed_cycles:
                    patient = cycle.patient
                    p_data = patient.patient_data
                    if not p_data: continue
                    
                    # Get wellness averages (simplified)
                    wellness_logs = WellnessLog.query.filter_by(user_id=patient.id).all()
                    if wellness_logs:
                        avg_stress = np.mean([l.stress_level for l in wellness_logs if l.stress_level])
                        avg_sleep = np.mean([l.sleep_hours for l in wellness_logs if l.sleep_hours])
                        avg_exercise = np.mean([l.exercise_minutes for l in wellness_logs if l.exercise_minutes])
                    else:
                        avg_stress, avg_sleep, avg_exercise = 3, 7.5, 30

                    features = np.array([[p_data.age or 30, p_data.bmi or 22, p_data.amh_level or 2.0, p_data.fsh_level or 6.0, p_data.previous_ivf_cycles or 0, avg_stress, avg_sleep, avg_exercise]])
                    prediction = ivf_model.predict(features)[0]
                    actual = 1 if cycle.outcome in ['BFP', 'Live Birth'] else 0
                    
                    if prediction == actual:
                        correct += 1
                    total += 1
                
                if total > 0:
                    model_stats['accuracy'] = round((correct / total) * 100, 1)
        except Exception as e:
            app.logger.error(f"Error calculating accuracy: {e}")

    return render_template('admin_dashboard.html', 
                           user=user, user_count=user_count, 
                           clinic_count=clinic_count, patient_count=patient_count,
                           doctor_count=doctor_count, model_stats=model_stats)

@app.route('/admin/clinics')
@admin_required
def admin_clinics():
    user = User.query.get(session['user_id'])
    clinics = Clinic.query.order_by(Clinic.name).all()
    return render_template('admin_clinics.html', user=user, clinics=clinics)

@app.route('/admin/clinic/add', methods=['GET', 'POST'])
@admin_required
def add_clinic():
    user = User.query.get(session['user_id'])
    if request.method == 'POST':
        new_clinic = Clinic(
            name=request.form['name'],
            address=request.form.get('address'),
            city=request.form.get('city'),
            state=request.form.get('state'),
            zip_code=request.form.get('zip_code'),
            phone=request.form.get('phone'),
            website=request.form.get('website'),
            latitude=float(request.form.get('latitude')) if request.form.get('latitude') else None,
            longitude=float(request.form.get('longitude')) if request.form.get('longitude') else None,
            description=request.form.get('description')
        )
        db.session.add(new_clinic)
        db.session.commit()
        flash('Clinic added successfully!', 'success')
        return redirect(url_for('admin_clinics'))
    return render_template('admin_clinic_form.html', user=user, clinic=None, title="Add New Clinic")

@app.route('/admin/clinic/edit/<int:clinic_id>', methods=['GET', 'POST'])
@admin_required
def edit_clinic(clinic_id):
    user = User.query.get(session['user_id'])
    clinic = Clinic.query.get_or_404(clinic_id)
    if request.method == 'POST':
        clinic.name = request.form['name']
        clinic.address = request.form.get('address')
        clinic.city = request.form.get('city')
        clinic.state = request.form.get('state')
        clinic.zip_code = request.form.get('zip_code')
        clinic.phone = request.form.get('phone')
        clinic.website = request.form.get('website')
        clinic.latitude = float(request.form.get('latitude')) if request.form.get('latitude') else None
        clinic.longitude = float(request.form.get('longitude')) if request.form.get('longitude') else None
        clinic.description = request.form.get('description')
        db.session.commit()
        flash('Clinic updated successfully!', 'success')
        return redirect(url_for('admin_clinics'))
    return render_template('admin_clinic_form.html', user=user, clinic=clinic, title="Edit Clinic")

@app.route('/admin/clinic/delete/<int:clinic_id>', methods=['POST'])
@admin_required
def delete_clinic(clinic_id):
    clinic = Clinic.query.get_or_404(clinic_id)
    # Optional: Check if any doctors are assigned to this clinic before deleting
    if clinic.doctors.first():
        flash('Cannot delete clinic. It has doctors assigned to it.', 'error')
        return redirect(url_for('admin_clinics'))
    
    db.session.delete(clinic)
    db.session.commit()
    flash('Clinic deleted successfully.', 'success')
    return redirect(url_for('admin_clinics'))

@app.route('/admin/users')
@admin_required
def admin_users():
    user = User.query.get(session['user_id'])
    users = User.query.order_by(User.created_at.desc()).all()
    return render_template('admin_users.html', user=user, users=users)

@app.route('/admin/user/edit/<int:user_id>', methods=['GET', 'POST'])
@admin_required
def edit_user(user_id):
    user_to_edit = User.query.get_or_404(user_id)
    user = User.query.get(session['user_id'])

    if request.method == 'POST':
        new_role = request.form.get('user_type')
        if new_role not in ['patient', 'doctor', 'admin']:
            flash('Invalid user role selected.', 'error')
            return redirect(url_for('admin_users'))
        
        user_to_edit.user_type = new_role
        db.session.commit()
        flash(f'User {user_to_edit.username}\'s role has been updated.', 'success')
        return redirect(url_for('admin_users'))

    return render_template('admin_user_form.html', user=user, user_to_edit=user_to_edit, title="Edit User Role")

@app.route('/admin/user/delete/<int:user_id>', methods=['POST'])
@admin_required
def delete_user(user_id):
    if user_id == session['user_id']:
        flash('You cannot delete your own account.', 'error')
        return redirect(url_for('admin_users'))
    user_to_delete = User.query.get_or_404(user_id)
    db.session.delete(user_to_delete)
    db.session.commit()
    flash(f'User {user_to_delete.username} has been deleted.', 'success')
    return redirect(url_for('admin_users'))

# IVF Success Predictor Page
@app.route('/ivf_predictor')
@login_required
def ivf_predictor():
    user = User.query.get(session['user_id'])
    patient_data = PatientData.query.filter_by(user_id=user.id).first()

    if not patient_data:
        flash('Please complete your medical profile first.', 'error')
        return redirect(url_for('update_profile'))

    # Calculate predictions using the service
    success_prediction = calculate_ivf_success_prediction(patient_data)
    embryo_quality_score = calculate_embryo_quality_score(patient_data)
    protocol_recommendations = generate_personalized_protocol(patient_data)

    # Save the new prediction to the database
    new_prediction = Prediction(
        user_id=user.id,
        success_probability=success_prediction.get('success_rate', 0) / 100.0, # Convert to 0-1.0 scale
        protocol_recommendation=protocol_recommendations.get('protocol_name'),
        llm_analysis=json.dumps({
            "success_factors": success_prediction.get('factors'),
            "embryo_score": embryo_quality_score.get('quality_score'),
            "protocol_optimizations": protocol_recommendations.get('success_optimization')
        })
    )
    db.session.add(new_prediction)
    db.session.commit()

    # Security: Disclaimer for AI usage
    disclaimer_text = "DISCLAIMER: This system uses Artificial Intelligence to provide estimates based on statistical data. It is NOT a diagnostic tool and should not replace professional medical advice. Always consult your fertility specialist."

    return render_template('ivf_predictor.html',
                          user=user,
                          patient_data=patient_data,
                          success_prediction=success_prediction,
                          embryo_quality_score=embryo_quality_score,
                          protocol_recommendations=protocol_recommendations,
                          disclaimer_text=disclaimer_text)

# Wellness tracking
@app.route('/wellness', methods=['GET', 'POST'])
@login_required
def wellness():
    user = User.query.get(session['user_id'])

    # Add a check to handle invalid sessions (e.g., after a database reset)
    if user is None:
        session.clear() # Clear the invalid session
        flash('Your session has expired. Please log in again.', 'error')
        return redirect(url_for('login'))

    if request.method == 'POST':
        # Safely parse date from form
        date_str = request.form.get('date')
        log_date = datetime.strptime(date_str, '%Y-%m-%d').date() if date_str else datetime.utcnow().date()
        
        # Check if a log for this date already exists
        existing_log = WellnessLog.query.filter_by(user_id=user.id, date=log_date).first()
        
        if existing_log:
            # Update the existing log
            wellness_log = existing_log
            flash_message = 'Wellness log updated successfully!'
        else:
            # Create a new log
            wellness_log = WellnessLog(user_id=user.id, date=log_date)
            db.session.add(wellness_log)
            flash_message = 'Wellness log saved successfully!'

        # Populate or update fields
        wellness_log.mood_rating = int(request.form.get('mood_rating', 0))
        wellness_log.mood_notes = request.form.get('mood_notes')

        # --- Analyze emotion from mood_notes and store it ---
        if wellness_log.mood_notes and emotion_model and vectorizer:
            text_to_analyze = wellness_log.mood_notes
            X = vectorizer.transform([text_to_analyze])
            wellness_log.detected_emotion = emotion_model.predict(X)[0]
        else:
            wellness_log.detected_emotion = None # Clear if no notes

        wellness_log.stress_level = int(request.form.get('stress_level', 0))
        wellness_log.stress_factors = request.form.get('stress_factors')
        wellness_log.sleep_hours = float(request.form.get('sleep_hours', 0) or 0)
        wellness_log.sleep_quality = int(request.form.get('sleep_quality', 0) or 0)
        wellness_log.sleep_notes = request.form.get('sleep_notes')
        wellness_log.symptoms = request.form.get('symptoms')
        wellness_log.energy_level = int(request.form.get('energy_level', 0))
        wellness_log.exercise_minutes = int(request.form.get('exercise_minutes', 0) or 0)
        wellness_log.meditation_minutes = int(request.form.get('meditation_minutes', 0) or 0)
        wellness_log.yoga_practiced = ('yoga_practiced' in request.form)
        wellness_log.water_intake = int(request.form.get('water_intake', 0) or 0)
        wellness_log.nutrition_score = int(request.form.get('nutrition_score', 0) or 0)
        wellness_log.supplements_taken = request.form.get('supplements_taken')

        # Save meal descriptions
        wellness_log.meal_breakfast = request.form.get('meal_breakfast')
        wellness_log.meal_lunch = request.form.get('meal_lunch')
        wellness_log.meal_dinner = request.form.get('meal_dinner')
        wellness_log.meal_snacks = request.form.get('meal_snacks')
        
        db.session.commit()

        # --- Smart Notification Triggers ---
        # Check last 3 wellness logs for negative trends
        last_3_logs = WellnessLog.query.filter_by(user_id=user.id).order_by(WellnessLog.date.desc()).limit(3).all()
        if len(last_3_logs) >= 3:
            all_low_mood = all((log.mood_rating or 5) <= 2 for log in last_3_logs)
            all_high_stress = all((log.stress_level or 1) >= 5 for log in last_3_logs)
            if all_low_mood:
                create_notification(
                    user_id=user.id,
                    message="Your mood has been low for 3 consecutive days. Consider speaking to your doctor or counselor.",
                    type='warning',
                    category='wellness',
                    link=url_for('wellness')
                )
            if all_high_stress:
                create_notification(
                    user_id=user.id,
                    message="High stress detected for 3 consecutive days. Try mindfulness exercises or consult your doctor.",
                    type='warning',
                    category='wellness',
                    link=url_for('mindfulness')
                )

        flash(flash_message, 'success')
        return redirect(url_for('wellness'))

    # Get recent logs for charts
    recent_logs = WellnessLog.query.filter_by(user_id=user.id).order_by(WellnessLog.date.desc()).limit(30).all()

    return render_template('wellness.html', user=user, recent_logs=recent_logs)

# Nutrition guidance
@app.route('/nutrition')
@login_required
def nutrition():
    user = User.query.get(session['user_id'])
    patient_data = PatientData.query.filter_by(user_id=user.id).first()

    regenerate = request.args.get('regenerate')

    # Get AI-generated guidance
    if regenerate in ['all', 'meal']:
        nutrition_plan = get_nutrition_plan(patient_data)
    else:
        nutrition_plan = get_nutrition_plan(patient_data)

    if regenerate in ['all', 'yoga']:
        yoga_routine = get_yoga_routine(patient_data)
    else:
        yoga_routine = get_yoga_routine(patient_data)

    # Get today's meal log and generate nutrition summary
    today_log = WellnessLog.query.filter_by(user_id=user.id, date=date.today()).first()
    nutrition_summary = {}

    if today_log and (today_log.meal_breakfast or today_log.meal_lunch or today_log.meal_dinner or today_log.meal_snacks):
        meal_descriptions = {
            "breakfast": today_log.meal_breakfast,
            "lunch": today_log.meal_lunch,
            "dinner": today_log.meal_dinner,
            "snacks": today_log.meal_snacks
        }
        # Remove empty meals before sending to AI
        meal_descriptions = {k: v for k, v in meal_descriptions.items() if v}

        if meal_descriptions:
            nutrition_summary = get_nutrition_analysis(meal_descriptions)

    return render_template('nutrition.html',
                          user=user,
                          nutrition_plan=nutrition_plan,
                          yoga_routine=yoga_routine,
                          today_log=today_log,
                          nutrition_summary=nutrition_summary)

# Profile management
@app.route('/update_profile', methods=['GET', 'POST'])
@login_required
def update_profile():
    user = User.query.get(session['user_id'])

    # Add a check to handle invalid sessions (e.g., after a database reset)
    if user is None:
        session.clear() # Clear the invalid session
        flash('Your session has expired. Please log in again.', 'error')
        return redirect(url_for('login'))

    patient_data = PatientData.query.filter_by(user_id=user.id).first()
    
    if not patient_data:
        patient_data = PatientData(user_id=user.id)
        db.session.add(patient_data)
    
    if request.method == 'POST':
        # Update user data
        user.first_name = request.form['first_name']
        user.last_name = request.form['last_name']
        user.phone = request.form.get('phone')
        
        if request.form.get('date_of_birth'):
            user.date_of_birth = datetime.strptime(request.form['date_of_birth'], '%Y-%m-%d').date()
        
        # Update patient data with type casting and validation
        if request.form.get('age'):
            patient_data.age = int(request.form['age'])
        
        # Safely handle float conversions
        height_cm = float(request.form.get('height', 0))
        weight_kg = float(request.form.get('weight', 0))
        
        if height_cm > 0:
            patient_data.height = height_cm
        if weight_kg > 0:
            patient_data.weight = weight_kg
            
        # Calculate BMI
        if patient_data.height and patient_data.height > 0 and patient_data.weight:
            patient_data.bmi = patient_data.weight / ((patient_data.height / 100) ** 2)
        else:
            patient_data.bmi = None 
        
        if request.form.get('amh_level'):
            patient_data.amh_level = float(request.form['amh_level'])
        if request.form.get('fsh_level'):
            patient_data.fsh_level = float(request.form['fsh_level'])
            
        patient_data.diagnosis = request.form.get('diagnosis')
        patient_data.medical_history = request.form.get('medical_history')
        patient_data.medications = request.form.get('medications')
        patient_data.allergies = request.form.get('allergies')
        patient_data.lifestyle_factors = request.form.get('lifestyle_factors')
        
        if request.form.get('previous_pregnancies'):
            patient_data.previous_pregnancies = int(request.form['previous_pregnancies'])
        if request.form.get('previous_ivf_cycles'):
            patient_data.previous_ivf_cycles = int(request.form['previous_ivf_cycles'])
        if request.form.get('partner_age'):
            patient_data.partner_age = int(request.form['partner_age'])
            
        patient_data.partner_diagnosis = request.form.get('partner_diagnosis')
        
        db.session.commit()
        flash('Profile updated successfully!', 'success')
        return redirect(url_for('patient_dashboard'))
    
    return render_template('update_profile.html', user=user, patient_data=patient_data)

# File upload with AI extraction and prediction
@app.route('/upload_document', methods=['POST'])
@login_required
def upload_document():
    extracted_medical_data = None
    ai_prediction = None
    auto_saved = False
    
    if 'file' not in request.files:
        flash('No file selected.', 'error')
        return redirect(request.referrer)

    file = request.files['file']
    if file.filename == '':
        flash('No file selected.', 'error')
        return redirect(request.referrer)

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_')
        
        # Get file extension
        file_ext = os.path.splitext(filename)[1] if '.' in filename else ''
        
        # Truncate filename to avoid Windows MAX_PATH limit (260 chars)
        base_name = os.path.splitext(filename)[0][:50]
        filename = timestamp + base_name + file_ext

        # Ensure the uploads folder exists
        upload_dir = app.config['UPLOAD_FOLDER']
        os.makedirs(upload_dir, exist_ok=True)

        file_path = os.path.join(upload_dir, filename)
        file.save(file_path)

        # Extract text from the file
        extracted_text = ""
        try:
            if file.content_type == 'application/pdf':
                extracted_text = pdf_extract_text(file_path)
            elif file.content_type.startswith('image/'):
                image = Image.open(file_path)
                extracted_text = pytesseract.image_to_string(image)
            elif 'word' in file.content_type or filename.endswith('.docx'):
                try:
                    from docx import Document
                    doc = Document(file_path)
                    extracted_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                except:
                    pass
            else:
                # Try to read as text
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        extracted_text = f.read()
                except:
                    pass
        except Exception as e:
            app.logger.error(f"Error extracting text: {e}")
        
        # Extract medical parameters from text
        if extracted_text:
            extracted_medical_data = extract_medical_params(extracted_text)
            
            # Generate automatic insights
            medical_insights = generate_medical_insights(extracted_medical_data)
            
            # Try to auto-fill user profile
            user = User.query.get(session['user_id'])
            patient_data = PatientData.query.filter_by(user_id=user.id).first()
            
            if patient_data and extracted_medical_data:
                # Auto-fill extracted values into patient profile
                if extracted_medical_data.get('age') and not patient_data.age:
                    patient_data.age = extracted_medical_data['age']
                    auto_saved = True
                
                if extracted_medical_data.get('amh') and not patient_data.amh_level:
                    patient_data.amh_level = extracted_medical_data['amh']
                    auto_saved = True
                
                if extracted_medical_data.get('fsh') and not patient_data.fsh_level:
                    patient_data.fsh_level = extracted_medical_data['fsh']
                    auto_saved = True
                
                if auto_saved:
                    db.session.commit()
            
            # Run AI prediction if we have enough data
            if patient_data and extracted_medical_data.get('amh') and ivf_model:
                try:
                    # Create a temporary patient data for prediction
                    pred_age = extracted_medical_data.get('age') or (patient_data.age if patient_data.age else 30)
                    pred_bmi = patient_data.bmi if patient_data.bmi else 22
                    pred_amh = extracted_medical_data.get('amh', 2.0)
                    pred_fsh = extracted_medical_data.get('fsh', 6.0)
                    pred_prev_ivf = patient_data.previous_ivf_cycles if patient_data.previous_ivf_cycles else 0
                    
                    # Use wellness data if available
                    recent_wellness = WellnessLog.query.filter_by(user_id=user.id).order_by(WellnessLog.date.desc()).first()
                    pred_stress = recent_wellness.stress_level if recent_wellness and recent_wellness.stress_level else 3
                    pred_sleep = recent_wellness.sleep_hours if recent_wellness and recent_wellness.sleep_hours else 7
                    pred_exercise = recent_wellness.exercise_minutes if recent_wellness and recent_wellness.exercise_minutes else 30
                    
                    features = np.array([[pred_age, pred_bmi, pred_amh, pred_fsh, pred_prev_ivf, pred_stress, pred_sleep, pred_exercise/30.0]])
                    prediction = ivf_model.predict(features)[0]
                    probability = ivf_model.predict_proba(features)[0]
                    success_chance = probability[1] * 100
                    
                    ai_prediction = {
                        'success_probability': round(success_chance, 1),
                        'prediction': 'Likely Success' if prediction == 1 else 'Low Chance'
                    }
                    
                    # Save prediction to database
                    new_prediction = Prediction(
                        user_id=user.id,
                        success_probability=success_chance / 100.0,
                        protocol_recommendation='Auto-generated from document',
                        llm_analysis=json.dumps({
                            'source': 'document_upload',
                            'extracted_params': extracted_medical_data
                        })
                    )
                    db.session.add(new_prediction)
                    db.session.commit()
                except Exception as e:
                    app.logger.error(f"Error in AI prediction: {e}")
        
        # Save document with extracted text
        document = MedicalDocument(
            user_id=session['user_id'],
            filename=filename,
            original_filename=file.filename,
            file_type=file.content_type,
            file_size=os.path.getsize(file_path),
            description=request.form.get('description'),
            extracted_text=extracted_text[:5000] if extracted_text else None
        )

        db.session.add(document)
        db.session.commit()
        
        # Show success message with AI summary info
        flash('Document uploaded successfully!', 'success')
        
        # Pass extracted data to the template via session for display
        if extracted_medical_data:
            session['last_extracted_medical'] = extracted_medical_data
            session['last_medical_insights'] = medical_insights if extracted_text else []
            session['last_ai_prediction'] = ai_prediction
            session['last_auto_saved'] = auto_saved
    else:
        flash('Invalid file type.', 'error')

    return redirect(url_for('my_documents'))

# My Documents page
@app.route('/my_documents')
@login_required
def my_documents():
    user = User.query.get(session['user_id'])

    # Add a check to handle invalid sessions (e.g., after a database reset)
    if user is None:
        session.clear() # Clear the invalid session
        flash('Your session has expired. Please log in again.', 'error')
        return redirect(url_for('login'))

    if user.user_type != 'patient':
        flash('Access denied.', 'error')
        return redirect(url_for('index'))

    documents = MedicalDocument.query.filter_by(user_id=user.id).order_by(MedicalDocument.uploaded_at.desc()).all()
    return render_template('my_documents.html', user=user, documents=documents)

# Download document
@app.route('/download_document/<int:doc_id>')
@login_required
def download_document(doc_id):
    document = MedicalDocument.query.get_or_404(doc_id)
    if document.user_id != session['user_id']:
        flash('Access denied.', 'error')
        return redirect(url_for('my_documents'))

    upload_dir = app.config['UPLOAD_FOLDER']
    return send_from_directory(upload_dir, document.filename, as_attachment=True, download_name=document.original_filename)

# Analyze document
@app.route('/analyze_document/<int:doc_id>')
@login_required
def analyze_document(doc_id):
    document = MedicalDocument.query.get_or_404(doc_id)
    if document.user_id != session['user_id']:
        flash('Access denied.', 'error')
        return redirect(url_for('my_documents'))

    upload_dir = app.config['UPLOAD_FOLDER']
    file_path = os.path.join(upload_dir, document.filename)

    if not os.path.exists(file_path):
        flash('Document file not found.', 'error')
        return redirect(url_for('my_documents'))

    extracted_text = ""

    try:
        if document.file_type == 'application/pdf':
            # Extract text from PDF
            extracted_text = pdf_extract_text(file_path)
        elif document.file_type.startswith('image/'):
            # Extract text from image using OCR
            image = Image.open(file_path)
            extracted_text = pytesseract.image_to_string(image)
        elif document.file_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
            # Handle DOC and DOCX files
            try:
                from docx import Document
                doc = Document(file_path)
                extracted_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            except ImportError:
                extracted_text = "DOCX processing not available. Please install python-docx library."
            except Exception as e:
                extracted_text = f"Error processing Word document: {str(e)}"
        else:
            # For other files, try to read as text
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    extracted_text = f.read()
            except UnicodeDecodeError: 
                extracted_text = "This file appears to be binary and cannot be analyzed as text."
    except Exception as e:
        flash(f'Error analyzing document: {str(e)}', 'error')
        return redirect(url_for('my_documents'))

    # Store the extracted text in the document record
    document.extracted_text = extracted_text
    db.session.commit()

    return render_template('document_analysis.html', document=document, extracted_text=extracted_text)

# Add Doctor's Note to Cycle
@app.route('/add_cycle_note/<int:cycle_id>', methods=['POST'])
@login_required
def add_cycle_note(cycle_id):
    user = User.query.get(session['user_id'])
    if user.user_type != 'doctor':
        flash('Only doctors can add notes.', 'error')
        return redirect(url_for('patient_dashboard'))

    cycle = IVFCycle.query.get_or_404(cycle_id)
    note_content = request.form.get('note_content')

    if not note_content:
        flash('Note content cannot be empty.', 'error')
        return redirect(url_for('doctor_dashboard'))

    new_note = CycleNote(
        cycle_id=cycle.id,
        doctor_id=user.id,
        note_content=note_content
    )
    db.session.add(new_note)
    db.session.commit()
    flash('Note added successfully to the patient\'s cycle.', 'success')
    return redirect(url_for('doctor_dashboard'))

# Create New IVF Cycle
@app.route('/create_cycle', methods=['POST'])
@login_required
def create_cycle():
    user = User.query.get(session['user_id'])
    if user.user_type != 'doctor':
        flash('Only doctors can create cycles.', 'error')
        return redirect(url_for('index'))

    patient_id = request.form.get('patient_id')
    protocol = request.form.get('protocol', 'Antagonist')
    start_date_str = request.form.get('start_date')
    patient_notes = request.form.get('patient_notes')

    if not patient_id or not start_date_str:
        flash('Patient and start date are required.', 'error')
        return redirect(url_for('doctor_dashboard'))

    try:
        start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
    except ValueError:
        flash('Invalid date format.', 'error')
        return redirect(url_for('doctor_dashboard'))

    # Check if patient exists and is a patient
    patient = User.query.get(patient_id)
    if not patient or patient.user_type != 'patient':
        flash('Invalid patient selected.', 'error')
        return redirect(url_for('doctor_dashboard'))

    new_cycle = IVFCycle(
        patient_id=patient.id,
        protocol=protocol,
        start_date=start_date,
        patient_notes=patient_notes
    )
    db.session.add(new_cycle)
    db.session.commit()
    flash(f'New IVF cycle created for {patient.first_name} {patient.last_name}.', 'success')
    return redirect(url_for('doctor_dashboard'))

# Schedule Appointment
@app.route('/schedule_appointment', methods=['GET', 'POST'])
@login_required
def schedule_appointment():
    user = User.query.get(session['user_id'])
    if user.user_type != 'doctor':
        flash('Only doctors can schedule appointments.', 'error')
        return redirect(url_for('index'))

    if request.method == 'GET':
        # Show the appointment scheduling form
        patients = User.query.filter_by(user_type='patient').all()
        recent_activities = MedicalActivity.query.join(User, MedicalActivity.patient_id == User.id).order_by(MedicalActivity.performed_date.desc()).limit(10).all()
        return render_template('schedule_appointment.html', user=user, patients=patients, recent_activities=recent_activities)

    # Handle POST request
    patient_id = request.form.get('patient_id')
    appointment_type = request.form.get('appointment_type')
    appointment_date_str = request.form.get('appointment_date')
    appointment_time_str = request.form.get('appointment_time')
    notes = request.form.get('notes')

    if not all([patient_id, appointment_type, appointment_date_str, appointment_time_str]):
        flash('All fields are required.', 'error')
        return redirect(url_for('schedule_appointment'))

    try:
        appointment_date = datetime.strptime(appointment_date_str, '%Y-%m-%d').date()
        appointment_time = datetime.strptime(appointment_time_str, '%H:%M').time()
        performed_date = datetime.combine(appointment_date, appointment_time)
    except ValueError:
        flash('Invalid date or time format.', 'error')
        return redirect(url_for('schedule_appointment'))

    # Check if patient exists
    patient = User.query.get(patient_id)
    if not patient or patient.user_type != 'patient':
        flash('Invalid patient selected.', 'error')
        return redirect(url_for('schedule_appointment'))

    new_activity = MedicalActivity(
        patient_id=patient.id,
        activity_type=appointment_type,
        activity_name=f"{appointment_type.title()} Appointment",
        performed_date=performed_date,
        notes=notes
    )
    db.session.add(new_activity)
    db.session.commit()
    flash(f'Appointment scheduled for {patient.first_name} {patient.last_name} on {performed_date.strftime("%B %d, %Y at %I:%M %p")}.', 'success')
    return redirect(url_for('doctor_dashboard'))

# --- Doctor Messaging Routes ---
@app.route('/doctor/messages')
@login_required
def doctor_messages():
    if session.get('user_type') != 'doctor':
        flash('Access denied.', 'error')
        return redirect(url_for('index'))

    doctor_id = session['user_id']
    
    # Subquery to get the latest message time for each conversation
    subquery = db.session.query(
        db.func.greatest(Message.sender_id, Message.receiver_id).label('user1'),
        db.func.least(Message.sender_id, Message.receiver_id).label('user2'),
        db.func.max(Message.created_at).label('last_message_time')
    ).filter(db.or_(Message.sender_id == doctor_id, Message.receiver_id == doctor_id)).group_by('user1', 'user2').subquery()

    # Join to get the full message and patient details
    conversations = db.session.query(User, Message).join(
        subquery,
        db.and_(
            db.func.greatest(Message.sender_id, Message.receiver_id) == subquery.c.user1,
            db.func.least(Message.sender_id, Message.receiver_id) == subquery.c.user2,
            Message.created_at == subquery.c.last_message_time
        )
    ).join(User, db.case(
        (Message.sender_id == doctor_id, User.id == Message.receiver_id),
        else_=User.id == Message.sender_id
    )).order_by(subquery.c.last_message_time.desc()).all()

    return render_template('doctor_messages.html', conversations=conversations)

@app.route('/doctor/messages/<int:patient_id>', methods=['GET', 'POST'])
@login_required
def view_conversation(patient_id):
    if session.get('user_type') != 'doctor':
        flash('Access denied.', 'error')
        return redirect(url_for('index'))

    doctor_id = session['user_id']
    patient = User.query.get_or_404(patient_id)

    if request.method == 'POST':
        content = request.form.get('content')
        if content:
            new_message = Message(sender_id=doctor_id, receiver_id=patient_id, subject="Re: Conversation", content=content)
            db.session.add(new_message)
            db.session.commit()
            flash('Reply sent successfully!', 'success')
            return redirect(url_for('view_conversation', patient_id=patient_id))

    messages = Message.query.filter(
        db.or_(
            db.and_(Message.sender_id == doctor_id, Message.receiver_id == patient_id),
            db.and_(Message.sender_id == patient_id, Message.receiver_id == doctor_id)
        )
    ).order_by(Message.created_at.asc()).all()

    return render_template('view_conversation.html', messages=messages, patient=patient)

# Chat API using Gemini
@app.route('/api/chat', methods=['POST'])
@login_required
def chat_api():
    try:
        data = request.get_json()
        message = data.get('message')
        
        if not message:
            return jsonify({'error': 'No message provided'}), 400
        
        user = User.query.get(session['user_id']) # type: ignore
        patient_data = None
        if user.user_type == 'patient':
            patient_data = PatientData.query.filter_by(user_id=user.id).first()
        
        # Get AI response using Gemini
        response = get_chatbot_response_gemini(message, user, patient_data)
        
        # Save chat to database
        chat_message = ChatMessage(
            user_id=user.id,
            message=message,
            response=response
        )
        db.session.add(chat_message)
        db.session.commit()
        
        return jsonify({
            'response': response,
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        app.logger.error(f"Chat API error: {str(e)}")
        return jsonify({'error': 'Failed to process chat message'}), 500


# TTS API using Gemini
@app.route('/api/chat/tts', methods=['POST'])
@login_required
def tts_api():
    try:
        data = request.get_json()
        text = data.get('text')
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        # Generate TTS audio using Gemini
        audio_result = generate_tts_audio(text)
        
        if 'error' in audio_result:
            return jsonify({'error': audio_result['error']}), 500
        
        return jsonify({
            'audio_data': audio_result.get('audio_data'),
            'mime_type': audio_result.get('mime_type', 'audio/pcm')
        })
        
    except Exception as e:
        app.logger.error(f"TTS API error: {str(e)}")
        return jsonify({'error': 'Failed to generate audio'}), 500

# Chatbot page
@app.route('/chatbot')
@login_required
def chatbot():
    user = User.query.get(session['user_id'])
    return render_template('chatbot.html', user=user)

# FAQ page
@app.route('/faq')
def faq():
    return render_template('faq.html')

# Welcome API endpoint
@app.route('/welcome')
@login_required
def welcome():
    app.logger.info(f"Request received: {request.method} {request.path}")
    return jsonify({'message': 'Welcome to the IVF Journey API!'})

# Hello API endpoint
@app.route('/hello')
@login_required
def hello():
    app.logger.info(f"Request received: {request.method} {request.path}")
    return jsonify({'message': 'Hello from the IVF Journey API!'})

# API endpoint for wellness data
@app.route('/api/wellness_data')
@login_required
def wellness_data():
    user_id = session['user_id']
    logs = WellnessLog.query.filter_by(user_id=user_id).order_by(WellnessLog.date.desc()).limit(30).all()
    
    data = {
        'dates': [log.date.strftime('%Y-%m-%d') for log in reversed(logs)],
        'mood': [log.mood_rating or 0 for log in reversed(logs)],
        'stress': [log.stress_level or 0 for log in reversed(logs)],
        'sleep_hours': [log.sleep_hours or 0 for log in reversed(logs)],
        'sleep_quality': [log.sleep_quality or 0 for log in reversed(logs)],
        'energy': [log.energy_level or 0 for log in reversed(logs)],
        'detected_emotions': [log.detected_emotion for log in reversed(logs)],
        # Additional fields for new charts
        'exercise_minutes': [log.exercise_minutes or 0 for log in reversed(logs)],
        'meditation_minutes': [log.meditation_minutes or 0 for log in reversed(logs)],
        'water_intake': [log.water_intake or 0 for log in reversed(logs)]
    }
    
    return jsonify(data)

# Generate AI image
@app.route('/api/generate_image', methods=['POST'])
@login_required
def generate_image_api():
    try:
        data = request.get_json()
        prompt = data.get('prompt')
        
        if not prompt:
            return jsonify({'error': 'No prompt provided'}), 400
        
        # Get AI image URL
        image_url = generate_medical_image(prompt)
        return jsonify({'image_url': image_url})
        
    except Exception as e:
        app.logger.error(f"Image generation error: {str(e)}")
        return jsonify({'error': 'Failed to generate image'}), 500

# --- New API endpoint to trigger ML prediction and store it ---
@app.route("/api/predict_ivf_ml", methods=["POST"])
@login_required # Ensure only logged-in users can trigger this
def api_predict_ivf_ml():
    from predict import predict_and_store # Import here to prevent circular dependency

    user_id = session.get('user_id') # Get user_id from session for security
    if not user_id:
        return jsonify({"error": "User not authenticated"}), 401
    
    try:
        res = predict_and_store(int(user_id))
        # Fetch the newly created prediction from DB to return more details if needed
        # For now, just return the result from predict_and_store
        return jsonify({"status": "ok", "result": res})
    except ValueError as e: # Catch specific errors from build_feature_vector
        app.logger.error(f"Prediction error for user {user_id}: {str(e)}")
        return jsonify({"status": "error", "message": str(e)}), 400
    except Exception as e:
        app.logger.error(f"Prediction error for user {user_id}: {str(e)}")
        return jsonify({"status": "error", "message": "An internal error occurred during prediction."}), 500
# ============ IVF SUCCESS IMPROVEMENT SIMULATOR API ============

@app.route("/api/ivf_simulator", methods=["POST"])
@login_required
def ivf_simulator_api():
    """
    API endpoint for the IVF Success Improvement Simulator.
    
    Accepts patient features JSON and returns:
    - Current success probability (baseline)
    - One-by-one lifestyle improvement simulations
    - Combined optimization scenario
    - Scientific explanations for each recommendation
    
    Uses the existing pre-trained model — NEVER retrains.
    Stores simulation results in the database for future reference.
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400

        # Extract features from the request (supports both camelCase and snake_case)
        features_mapping = {
            'age': data.get('age') or data.get('Age'),
            'bmi': data.get('bmi') or data.get('BMI'),
            'amh': data.get('amh') or data.get('AMH'),
            'fsh': data.get('fsh') or data.get('FSH'),
            'previous_ivf': data.get('previous_ivf') or data.get('Previous_IVF_Attempts') or 0,
            'stress': data.get('stress') or data.get('Stress_Level'),
            'sleep_hours': data.get('sleep_hours') or data.get('Sleep_Hours'),
            'exercise_min': data.get('exercise_min') or data.get('Exercise_Min_per_Day')
        }

        # Convert to float where possible, remove None values
        clean_features = {}
        for k, v in features_mapping.items():
            if v is not None:
                try:
                    clean_features[k] = float(v)
                except (ValueError, TypeError):
                    pass

        # Import and run the simulator
        from prediction_service import calculate_ivf_improvements
        result = calculate_ivf_improvements(clean_features)

        # Check for errors
        if result.get('error'):
            return jsonify(result), 422

        # Store the simulation result in database for audit/reference
        try:
            user_id = session['user_id']
            sim_record = SimulationResult(
                user_id=user_id,
                input_features=json.dumps(clean_features),
                current_probability=result['current_probability'],
                simulation_results=json.dumps(result['improvements'])
            )
            db.session.add(sim_record)
            db.session.commit()
            result['simulation_id'] = sim_record.id
        except Exception as e:
            app.logger.error(f"Failed to save simulation result: {e}")
            db.session.rollback()
            # Still return the results even if DB save fails
            result['simulation_id'] = None

        return jsonify(result)

    except Exception as e:
        app.logger.error(f"IVF Simulator API error: {str(e)}")
        return jsonify({
            'error': 'Failed to run improvement simulation',
            'current_probability': None,
            'improvements': []
        }), 500


# --- ML Model API Endpoints ---

@app.route("/predict_ivf", methods=["POST"])
def predict_ivf():
    """
    Predict IVF success using the new prediction service.
    Falls back to rule-based predictions if ML models are not available.
    """
    from prediction_service import calculate_ivf_success_prediction
    
    data = request.json
    app.logger.debug(f"Received /predict_ivf request data: {data}")
    
    try:
        # Create a simple object with the required attributes for prediction_service
        class PatientData:
            def __init__(self, data):
                self.age = data.get("Age")
                self.bmi = data.get("BMI")
                self.amh_level = data.get("AMH")
                self.fsh_level = data.get("FSH")
                self.previous_ivf_cycles = data.get("Previous_IVF_Attempts")
                self.stress = data.get("Stress_Level")
                self.sleep_hours = data.get("Sleep_Hours")
                self.exercise_min = data.get("Exercise_Min_per_Day")
                # Additional fields for better predictions
                self.partner_age = None
                self.lifestyle_factors = None
        
        patient_data = PatientData(data)
        
        # Use the prediction service (has ML + fallback to rule-based)
        result = calculate_ivf_success_prediction(patient_data)
        
        # Return in the format expected by the JavaScript
        return jsonify({
            "prediction_text": result.get("prediction_label", "Unknown"),
            "success_probability": result.get("success_rate", 35.0),
            "confidence": result.get("confidence", 70.0),
            "model_type": result.get("model_type", "unknown"),
            "interpretation": result.get("interpretation", "")
        })
        
    except Exception as e:
        app.logger.error(f"Prediction error: {str(e)}")
        # Fallback to basic response if everything fails
        return jsonify({
            "prediction_text": "Unable to predict",
            "success_probability": 35.0,
            "confidence": 50.0,
            "model_type": "error_fallback",
            "error": str(e)
        })

# endpoint for generating diagrams (client-side prompt)
@app.route("/generate_diagram", methods=["POST"])
@login_required
def api_generate_diagram():
    data = request.get_json()
    prompt = data.get("prompt")
    if not prompt:
        return jsonify({"error": "No prompt provided"}), 400
    try:
        url = generate_diagram(prompt)
        return jsonify({"url": url})
    except Exception as e:
        app.logger.error(f"Diagram generation failed: {e}")
        return jsonify({"error": "Failed to generate diagram"}), 500

@app.route("/predict_mood", methods=["POST"])
def predict_mood():
    if not mood_model:
        return jsonify({"error": "Mood prediction model not loaded."}), 503

    data = request.json
    try:
        X = np.array([[data["Mood_Rating"], data["Stress_Level"], data["Sleep_Hours"], data["Steps"]]])
        next_mood = mood_model.predict(X)[0]
        return jsonify({"Predicted_Next_Day_Mood": round(float(next_mood), 2)})
    except KeyError as e:
        return jsonify({"error": f"Missing feature: {e}"}), 400
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/analyze_emotion", methods=["POST"])
def analyze_emotion():
    if not emotion_model or not vectorizer:
        return jsonify({"error": "Emotion analysis model not loaded."}), 503

    text = request.json.get("text")
    if not text:
        return jsonify({"error": "Text not provided"}), 400
    X = vectorizer.transform([text])
    emotion = emotion_model.predict(X)[0]
    return jsonify({"Detected_Emotion": emotion})

# --- New API endpoint for mood analysis ---
@app.route('/api/analyze_mood', methods=['POST'])
@login_required
def analyze_mood():
    if not emotion_model or not vectorizer:
        return jsonify({'error': 'Emotion analysis model not loaded.'}), 503

    data = request.get_json()
    text = data.get('text')

    if not text:
        return jsonify({'error': 'No text provided'}), 400

    X = vectorizer.transform([text])
    emotion = emotion_model.predict(X)[0]

    # Return the detected emotion
    return jsonify({'emotion': emotion})



# --- My Reminders Route ---
@app.route('/my_reminders', methods=['GET', 'POST'])
@login_required
def my_reminders():
    user = User.query.get(session['user_id'])
    
    if user.user_type != 'patient':
        flash('Access denied.', 'error')
        return redirect(url_for('index'))
    
    if request.method == 'POST':
        # Add new medication reminder
        medication_name = request.form.get('medication_name')
        dosage = request.form.get('dosage')
        frequency = request.form.get('frequency')
        time_str = request.form.get('time_of_day')
        start_date_str = request.form.get('start_date')
        end_date_str = request.form.get('end_date')
        
        if not medication_name or not start_date_str:
            flash('Medication name and start date are required.', 'error')
            return redirect(url_for('my_reminders'))
        
        try:
            start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
            end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date() if end_date_str else None
            
            time_of_day = None
            if time_str:
                time_obj = datetime.strptime(time_str, '%H:%M').time()
                time_of_day = time_obj
            
            reminder = MedicationReminder(
                user_id=user.id,
                medication_name=medication_name,
                dosage=dosage,
                frequency=frequency,
                time_of_day=time_of_day,
                start_date=start_date,
                end_date=end_date,
                is_active=True
            )
            db.session.add(reminder)
            db.session.commit()
            flash('Medication reminder added successfully!', 'success')
        except Exception as e:
            flash(f'Error adding reminder: {str(e)}', 'error')
        
        return redirect(url_for('my_reminders'))
    
    # Get all reminders for this user
    reminders = MedicationReminder.query.filter_by(user_id=user.id).order_by(MedicationReminder.time_of_day).all()
    active_reminders = [r for r in reminders if r.is_active]
    inactive_reminders = [r for r in reminders if not r.is_active]
    
    return render_template('my_reminders.html', 
                          user=user, 
                          active_reminders=active_reminders,
                          inactive_reminders=inactive_reminders)


# --- Toggle Reminder Status ---
@app.route('/toggle_reminder/<int:reminder_id>', methods=['POST'])
@login_required
def toggle_reminder(reminder_id):
    user = User.query.get(session['user_id'])
    reminder = MedicationReminder.query.get_or_404(reminder_id)
    
    if reminder.user_id != user.id:
        flash('Access denied.', 'error')
        return redirect(url_for('my_reminders'))
    
    reminder.is_active = not reminder.is_active
    db.session.commit()
    flash(f'Reminder {"activated" if reminder.is_active else "deactivated"} successfully!', 'success')
    return redirect(url_for('my_reminders'))


# --- Delete Reminder ---
@app.route('/delete_reminder/<int:reminder_id>', methods=['POST'])
@login_required
def delete_reminder(reminder_id):
    user = User.query.get(session['user_id'])
    reminder = MedicationReminder.query.get_or_404(reminder_id)
    
    if reminder.user_id != user.id:
        flash('Access denied.', 'error')
        return redirect(url_for('my_reminders'))
    
    db.session.delete(reminder)
    db.session.commit()
    flash('Reminder deleted successfully!', 'success')
    return redirect(url_for('my_reminders'))


# --- My Treatment Route ---
@app.route('/my_treatment', methods=['GET', 'POST'])
@login_required
def my_treatment():
    user = User.query.get(session['user_id'])
    
    if user.user_type != 'patient':
        flash('Access denied.', 'error')
        return redirect(url_for('index'))
    
    if request.method == 'POST':
        # Add new IVF cycle (patient can add their own)
        protocol = request.form.get('protocol', 'Not Specified')
        start_date_str = request.form.get('start_date')
        patient_notes = request.form.get('patient_notes')
        
        if not start_date_str:
            flash('Start date is required.', 'error')
            return redirect(url_for('my_treatment'))
        
        try:
            start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
            
            # Check if there's already an active cycle
            active_cycle = IVFCycle.query.filter_by(patient_id=user.id).filter(
                IVFCycle.status.in_(['stimulation', 'retrieval', 'transfer', 'wait'])
            ).first()
            
            if active_cycle:
                flash('You already have an active treatment cycle. Please complete or update it first.', 'warning')
                return redirect(url_for('my_treatment'))
            
            cycle = IVFCycle(
                patient_id=user.id,
                protocol=protocol,
                start_date=start_date,
                patient_notes=patient_notes,
                status='stimulation'
            )
            db.session.add(cycle)
            db.session.commit()
            flash('Treatment cycle added successfully!', 'success')
        except Exception as e:
            flash(f'Error adding treatment cycle: {str(e)}', 'error')
        
        return redirect(url_for('my_treatment'))
    
    # Get all cycles for this user
    cycles = IVFCycle.query.filter_by(patient_id=user.id).order_by(IVFCycle.start_date.desc()).all()
    
    # Get medication history from medical activities
    medications = MedicalActivity.query.filter_by(patient_id=user.id, activity_type='injection').order_by(MedicalActivity.performed_date.desc()).all()
    
    return render_template('my_treatment.html', 
                          user=user, 
                          cycles=cycles,
                          medications=medications)


# --- Update Treatment Cycle Status ---
@app.route('/update_cycle_status/<int:cycle_id>', methods=['POST'])
@login_required
def update_cycle_status(cycle_id):
    user = User.query.get(session['user_id'])
    cycle = IVFCycle.query.get_or_404(cycle_id)
    
    if cycle.patient_id != user.id:
        flash('Access denied.', 'error')
        return redirect(url_for('my_treatment'))
    
    new_status = request.form.get('status')
    if new_status:
        cycle.status = new_status
        db.session.commit()
        flash('Cycle status updated successfully!', 'success')
    
    return redirect(url_for('my_treatment'))


# --- Messages Route ---
@app.route('/messages', methods=['GET', 'POST'])
@login_required
def messages():
    user = User.query.get(session['user_id'])
    
    if user.user_type != 'patient':
        flash('Access denied.', 'error')
        return redirect(url_for('index'))
    
    if request.method == 'POST':
        # Send a new message to a doctor
        receiver_id = request.form.get('receiver_id')
        subject = request.form.get('subject')
        content = request.form.get('content')
        
        if not receiver_id or not subject or not content:
            flash('All fields are required.', 'error')
            return redirect(url_for('messages'))
        
        # Verify the receiver is a doctor
        receiver = User.query.get(receiver_id)
        if not receiver or receiver.user_type != 'doctor':
            flash('Invalid doctor selected.', 'error')
            return redirect(url_for('messages'))
        
        message = Message(
            sender_id=user.id,
            receiver_id=receiver_id,
            subject=subject,
            content=content
        )
        db.session.add(message)
        db.session.commit()
        flash('Message sent successfully!', 'success')
        return redirect(url_for('messages'))
    
    # Get all doctors for the dropdown
    doctors = User.query.filter_by(user_type='doctor').all()
    
    # Get sent messages
    sent_messages = Message.query.filter_by(sender_id=user.id).order_by(Message.created_at.desc()).all()
    
    # Get received messages
    received_messages = Message.query.filter_by(receiver_id=user.id).order_by(Message.created_at.desc()).all()
    
    # Get unread count
    unread_count = Message.query.filter_by(receiver_id=user.id, is_read=False).count()
    
    return render_template('messages.html',
                          user=user,
                          doctors=doctors,
                          sent_messages=sent_messages,
                          received_messages=received_messages,
                          unread_count=unread_count)


# --- Mark Message as Read ---
@app.route('/mark_message_read/<int:message_id>', methods=['POST'])
@login_required
def mark_message_read(message_id):
    user = User.query.get(session['user_id'])
    message = Message.query.get_or_404(message_id)
    
    if message.receiver_id != user.id:
        flash('Access denied.', 'error')
        return redirect(url_for('messages'))
    
    message.is_read = True
    db.session.commit()
    return redirect(url_for('messages'))


# --- Mindfulness Route ---
@app.route('/mindfulness')
@login_required
def mindfulness():
    user = User.query.get(session['user_id'])
    
    # Mindfulness content - static for now
    mindfulness_exercises = [
        {
            'title': 'Deep Breathing',
            'description': 'A simple breathing exercise to reduce stress and anxiety.',
            'duration': '5 minutes',
            'category': 'Breathing'
        },
        {
            'title': 'Progressive Muscle Relaxation',
            'description': 'Release tension by systematically relaxing each muscle group.',
            'duration': '15 minutes',
            'category': 'Relaxation'
        },
        {
            'title': 'Guided Visualization',
            'description': 'Imagine a peaceful place to promote calm and relaxation.',
            'duration': '10 minutes',
            'category': 'Visualization'
        },
        {
            'title': 'Body Scan Meditation',
            'description': 'Bring awareness to each part of your body to release tension.',
            'duration': '20 minutes',
            'category': 'Meditation'
        },
        {
            'title': 'Gratitude Journaling',
            'description': 'Write down things you are grateful for to shift focus to positivity.',
            'duration': '10 minutes',
            'category': 'Journaling'
        },
        {
            'title': 'Mindful Walking',
            'description': 'Walk slowly and mindfully, focusing on each step and sensation.',
            'duration': '15 minutes',
            'category': 'Movement'
        }
    ]
    
    ivf_specific_tips = [
        'Remember to take one day at a time during your IVF journey.',
        'It\'s normal to feel a range of emotions. Be kind to yourself.',
        'Connecting with others who understand can be very helpful.',
        'Taking time for self-care is important, not selfish.',
        'Every step forward is progress, no matter how small.'
    ]
    
    return render_template('mindfulness.html',
                          user=user,
                          exercises=mindfulness_exercises,
                          tips=ivf_specific_tips)


# --- API for Reminder Alerts ---
@app.route('/api/reminder_alerts')
@login_required
def reminder_alerts():
    """API endpoint to get active medication reminders"""
    user = User.query.get(session['user_id'])
    
    if user.user_type != 'patient':
        return jsonify({'error': 'Access denied'}), 403
    
    # Get active reminders
    active_reminders = MedicationReminder.query.filter_by(user_id=user.id, is_active=True).all()
    
    reminders_data = []
    for reminder in active_reminders:
        reminders_data.append({
            'id': reminder.id,
            'medication_name': reminder.medication_name,
            'dosage': reminder.dosage,
            'frequency': reminder.frequency,
            'time_of_day': reminder.time_of_day.strftime('%H:%M') if reminder.time_of_day else None
        })
    
    return jsonify({'reminders': reminders_data, 'count': len(reminders_data)})


# --- Notification API Routes ---
@app.route('/api/notifications')
@login_required
def api_notifications():
    """Fetch notifications for the current user."""
    user = User.query.get(session['user_id'])
    limit = request.args.get('limit', 20, type=int)
    unread_only = request.args.get('unread_only', 'false').lower() == 'true'
    
    query = Notification.query.filter_by(user_id=user.id)
    if unread_only:
        query = query.filter_by(is_read=False)
    
    notifications = query.order_by(Notification.created_at.desc()).limit(limit).all()
    
    data = []
    for n in notifications:
        data.append({
            'id': n.id,
            'message': n.message,
            'type': n.type,
            'category': n.category,
            'is_read': n.is_read,
            'link': n.link,
            'created_at': n.created_at.isoformat()
        })
    
    unread_count = Notification.query.filter_by(user_id=user.id, is_read=False).count()
    return jsonify({'notifications': data, 'unread_count': unread_count})


@app.route('/api/notifications/<int:notification_id>/read', methods=['POST'])
@login_required
def mark_notification_read(notification_id):
    """Mark a single notification as read."""
    user = User.query.get(session['user_id'])
    notification = Notification.query.get_or_404(notification_id)
    
    if notification.user_id != user.id:
        return jsonify({'error': 'Access denied'}), 403
    
    notification.is_read = True
    db.session.commit()
    return jsonify({'status': 'ok'})


@app.route('/api/notifications/read_all', methods=['POST'])
@login_required
def mark_all_notifications_read():
    """Mark all notifications for the user as read."""
    user = User.query.get(session['user_id'])
    Notification.query.filter_by(user_id=user.id, is_read=False).update({'is_read': True})
    db.session.commit()
    return jsonify({'status': 'ok', 'message': 'All notifications marked as read'})


@app.route('/notifications')
@login_required
def notifications_page():
    """Page to view all notifications."""
    user = User.query.get(session['user_id'])
    notifications = Notification.query.filter_by(user_id=user.id).order_by(Notification.created_at.desc()).all()
    return render_template('notifications.html', user=user, notifications=notifications)


# --- Partner Sync Routes ---
@app.route('/invite/partner', methods=['GET', 'POST'])
@login_required
def invite_partner():
    user = User.query.get(session['user_id'])
    if user is None or user.user_type != 'patient':
        flash('Only patients can invite a partner.', 'error')
        return redirect(url_for('patient_dashboard'))

    if request.method == 'POST':
        partner_name = request.form.get('partner_name', '').strip()
        partner_email = request.form.get('partner_email', '').strip().lower()
        phone = request.form.get('phone', '').strip()
        relationship = request.form.get('relationship', 'Partner')

        if not partner_name or not partner_email:
            flash('Partner name and email are required.', 'error')
            return redirect(url_for('invite_partner'))

        token = secrets.token_urlsafe(32)
        invitation = PartnerInvitation(
            patient_id=user.id,
            partner_email=partner_email,
            partner_name=partner_name,
            relationship=relationship,
            token_hash=hash_invitation_token(token),
            status='pending',
            expires_at=datetime.utcnow() + timedelta(days=7),
        )
        db.session.add(invitation)
        db.session.commit()

        accept_url = url_for('accept_partner_invitation', token=token, _external=True)
        flash(f'Invitation sent to {partner_email}. Share this link with your partner: {accept_url}', 'success')
        log_action(user.id, 'PARTNER_INVITATION_SENT', f'Invitation sent to {partner_email}')
        return redirect(url_for('patient_dashboard'))

    return render_template('partner_invite.html', user=user)


@app.route('/invite/partner/accept/<token>', methods=['GET', 'POST'])
def accept_partner_invitation(token):
    invitation = find_valid_invitation(token)
    if invitation is None:
        flash('This invitation is invalid, expired, or has already been processed.', 'error')
        return redirect(url_for('login'))

    if request.method == 'POST':
        action = request.form.get('action', 'accept')
        if action == 'decline':
            invitation.status = 'declined'
            invitation.declined_at = datetime.utcnow()
            db.session.commit()
            flash('Invitation declined. You can still request a new invitation later.', 'info')
            return redirect(url_for('login'))

        user = User.query.filter_by(email=invitation.partner_email.lower()).first()
        if user is None:
            return redirect(url_for('register_partner', token=token))
        if user.user_type != 'partner':
            flash('This invitation is for a partner account.', 'error')
            return redirect(url_for('login'))

        if user.id == invitation.patient_id:
            flash('This invitation is not for the current account.', 'error')
            return redirect(url_for('login'))

        connection = PartnerConnection.query.filter_by(patient_id=invitation.patient_id, partner_id=user.id).first()
        if connection is None:
            connection = PartnerConnection(patient_id=invitation.patient_id, partner_id=user.id, status='active', accepted_at=datetime.utcnow())
            db.session.add(connection)
            db.session.commit()
        else:
            connection.status = 'active'
            connection.accepted_at = datetime.utcnow()
            db.session.commit()

        permissions = PartnerSharingPermission.query.filter_by(connection_id=connection.id).first()
        if permissions is None:
            permissions = PartnerSharingPermission(
                connection_id=connection.id,
                treatment_shared=True,
                appointments_shared=True,
                medications_shared=True,
                wellness_shared=False,
                mood_shared=False,
                nutrition_shared=True,
                documents_shared=False,
                doctor_notes_shared=False,
                messages_enabled=True,
                tasks_enabled=True,
            )
            db.session.add(permissions)
        invitation.status = 'accepted'
        invitation.accepted_at = datetime.utcnow()
        db.session.commit()

        session['user_id'] = user.id
        session['user_type'] = 'partner'
        flash('Invitation accepted. You are now connected to your partner.', 'success')
        return redirect(url_for('partner_dashboard'))

    current_user = User.query.get(session.get('user_id'))
    return render_template('partner_invitation_accept.html', invitation=invitation, current_user=current_user)


@app.route('/invite/partner/decline/<token>', methods=['POST'])
def decline_partner_invitation(token):
    invitation = find_valid_invitation(token)
    if invitation is None:
        flash('This invitation is invalid or already processed.', 'error')
        return redirect(url_for('login'))
    invitation.status = 'declined'
    invitation.declined_at = datetime.utcnow()
    db.session.commit()
    flash('Invitation declined.', 'info')
    return redirect(url_for('login'))


@app.route('/partner/dashboard')
@login_required
def partner_dashboard():
    user = User.query.get(session['user_id'])
    if user is None or user.user_type != 'partner':
        flash('Access denied. This dashboard is for partners only.', 'error')
        return redirect(url_for('index'))

    connection = PartnerConnection.query.filter_by(partner_id=user.id, status='active').first()
    if not connection:
        flash('You are not connected to a patient yet. Please accept an invitation first.', 'info')
        return render_template('partner_dashboard.html', user=user, connection=None, patient=None, permissions=None)

    patient = User.query.get(connection.patient_id)
    permissions = get_partner_permissions(connection)
    latest_cycle = IVFCycle.query.filter_by(patient_id=patient.id).order_by(IVFCycle.start_date.desc()).first()
    latest_prediction = Prediction.query.filter_by(user_id=patient.id).order_by(Prediction.prediction_date.desc()).first()
    today_tasks = SharedTask.query.filter_by(connection_id=connection.id).count()
    pending_tasks = SharedTask.query.filter_by(connection_id=connection.id, status='pending').count()
    completed_tasks = SharedTask.query.filter_by(connection_id=connection.id, status='completed').count()
    wellness_logs = WellnessLog.query.filter_by(user_id=patient.id).order_by(WellnessLog.date.desc()).limit(7).all()
    reminders = MedicationReminder.query.filter_by(user_id=patient.id, is_active=True).count()
    checkins = PartnerCheckIn.query.filter_by(connection_id=connection.id).order_by(PartnerCheckIn.created_at.desc()).limit(5).all()
    notifications = PartnerNotification.query.filter_by(user_id=user.id).order_by(PartnerNotification.created_at.desc()).limit(5).all()
    return render_template(
        'partner_dashboard.html',
        user=user,
        patient=patient,
        connection=connection,
        permissions=permissions,
        latest_cycle=latest_cycle,
        latest_prediction=latest_prediction,
        today_tasks=today_tasks,
        pending_tasks=pending_tasks,
        completed_tasks=completed_tasks,
        wellness_logs=wellness_logs,
        reminders=reminders,
        checkins=checkins,
        notifications=notifications,
    )


@app.route('/partner/journey')
@login_required
def partner_journey():
    user = User.query.get(session['user_id'])
    if user.user_type != 'partner':
        flash('Access denied.', 'error')
        return redirect(url_for('index'))
    connection = PartnerConnection.query.filter_by(partner_id=user.id, status='active').first()
    if not connection:
        flash('No active partner connection found.', 'error')
        return redirect(url_for('partner_dashboard'))
    patient = User.query.get(connection.patient_id)
    cycles = IVFCycle.query.filter_by(patient_id=patient.id).order_by(IVFCycle.start_date.desc()).all()
    permissions = get_partner_permissions(connection)
    return render_template('partner_journey.html', user=user, patient=patient, cycles=cycles, connection=connection, permissions=permissions)


@app.route('/partner/appointments')
@login_required
def partner_appointments():
    user = User.query.get(session['user_id'])
    if user.user_type != 'partner':
        flash('Access denied.', 'error')
        return redirect(url_for('index'))
    connection = PartnerConnection.query.filter_by(partner_id=user.id, status='active').first()
    if not connection:
        flash('No active partner connection found.', 'error')
        return redirect(url_for('partner_dashboard'))
    permissions = get_partner_permissions(connection)
    if not permissions.appointments_shared:
        flash('Appointments are currently hidden for privacy.', 'info')
        return redirect(url_for('partner_dashboard'))
    patient = User.query.get(connection.patient_id)
    appointments = MedicalActivity.query.filter_by(patient_id=patient.id).order_by(MedicalActivity.performed_date.desc()).all()
    return render_template('partner_appointments.html', user=user, patient=patient, appointments=appointments, connection=connection, permissions=permissions)


@app.route('/partner/medications')
@login_required
def partner_medications():
    user = User.query.get(session['user_id'])
    if user.user_type != 'partner':
        flash('Access denied.', 'error')
        return redirect(url_for('index'))
    connection = PartnerConnection.query.filter_by(partner_id=user.id, status='active').first()
    if not connection:
        flash('No active partner connection found.', 'error')
        return redirect(url_for('partner_dashboard'))
    permissions = get_partner_permissions(connection)
    if not permissions.medications_shared:
        flash('Medication details are currently hidden.', 'info')
        return redirect(url_for('partner_dashboard'))
    patient = User.query.get(connection.patient_id)
    reminders = MedicationReminder.query.filter_by(user_id=patient.id, is_active=True).order_by(MedicationReminder.time_of_day).all()
    return render_template('partner_medications.html', user=user, patient=patient, reminders=reminders, connection=connection, permissions=permissions)


@app.route('/partner/wellness')
@login_required
def partner_wellness():
    user = User.query.get(session['user_id'])
    if user.user_type != 'partner':
        flash('Access denied.', 'error')
        return redirect(url_for('index'))
    connection = PartnerConnection.query.filter_by(partner_id=user.id, status='active').first()
    if not connection:
        flash('No active partner connection found.', 'error')
        return redirect(url_for('partner_dashboard'))
    permissions = get_partner_permissions(connection)
    if not permissions.wellness_shared:
        flash('Wellness sharing is currently disabled.', 'info')
        return redirect(url_for('partner_dashboard'))
    patient = User.query.get(connection.patient_id)
    wellness_logs = WellnessLog.query.filter_by(user_id=patient.id).order_by(WellnessLog.date.desc()).limit(14).all()
    return render_template('partner_wellness.html', user=user, patient=patient, wellness_logs=wellness_logs, connection=connection, permissions=permissions)


@app.route('/partner/tasks', methods=['GET', 'POST'])
@login_required
def partner_tasks():
    user = User.query.get(session['user_id'])
    if user.user_type != 'partner':
        flash('Access denied.', 'error')
        return redirect(url_for('index'))
    connection = PartnerConnection.query.filter_by(partner_id=user.id, status='active').first()
    if not connection:
        flash('No active partner connection found.', 'error')
        return redirect(url_for('partner_dashboard'))
    permissions = get_partner_permissions(connection)
    if not permissions.tasks_enabled:
        flash('Shared tasks are currently disabled.', 'info')
        return redirect(url_for('partner_dashboard'))

    if request.method == 'POST':
        title = request.form.get('title', '').strip()
        description = request.form.get('description', '').strip()
        due_date = request.form.get('due_date')
        priority = request.form.get('priority', 'medium')
        if title:
            task = SharedTask(
                connection_id=connection.id,
                title=title,
                description=description,
                assigned_to='partner',
                due_date=datetime.strptime(due_date, '%Y-%m-%d').date() if due_date else None,
                priority=priority,
                status='pending',
                created_by=user.id)
            db.session.add(task)
            db.session.commit()
            flash('Shared task added.', 'success')
            return redirect(url_for('partner_tasks'))

    tasks = SharedTask.query.filter_by(connection_id=connection.id).order_by(SharedTask.due_date.is_(None), SharedTask.due_date.asc(), SharedTask.created_at.desc()).all()
    return render_template('partner_tasks.html', user=user, connection=connection, tasks=tasks, permissions=permissions)


@app.route('/partner/tasks/<int:task_id>/complete', methods=['POST'])
@login_required
def complete_partner_task(task_id):
    user = User.query.get(session['user_id'])
    if user.user_type != 'partner':
        flash('Access denied.', 'error')
        return redirect(url_for('index'))
    task = SharedTask.query.get_or_404(task_id)
    connection = PartnerConnection.query.filter_by(partner_id=user.id, status='active').first()
    if task.connection_id != connection.id:
        flash('Access denied.', 'error')
        return redirect(url_for('partner_tasks'))
    task.status = 'completed'
    task.completed_at = datetime.utcnow()
    db.session.commit()
    flash('Task marked as completed.', 'success')
    return redirect(url_for('partner_tasks'))


@app.route('/partner/notes', methods=['GET', 'POST'])
@login_required
def partner_notes():
    user = User.query.get(session['user_id'])
    if user.user_type != 'partner':
        flash('Access denied.', 'error')
        return redirect(url_for('index'))
    connection = PartnerConnection.query.filter_by(partner_id=user.id, status='active').first()
    if not connection:
        flash('No active partner connection found.', 'error')
        return redirect(url_for('partner_dashboard'))

    if request.method == 'POST':
        title = request.form.get('title', '').strip()
        description = request.form.get('description', '').strip()
        if title and description:
            note = SharedNote(connection_id=connection.id, title=title, description=description, created_by=user.id)
            db.session.add(note)
            db.session.commit()
            flash('Shared note added.', 'success')
            return redirect(url_for('partner_notes'))

    notes = SharedNote.query.filter_by(connection_id=connection.id).order_by(SharedNote.created_at.desc()).all()
    return render_template('partner_notes.html', user=user, connection=connection, notes=notes)


@app.route('/partner/messages', methods=['GET', 'POST'])
@login_required
def partner_messages():
    user = User.query.get(session['user_id'])
    if user.user_type != 'partner':
        flash('Access denied.', 'error')
        return redirect(url_for('index'))
    connection = PartnerConnection.query.filter_by(partner_id=user.id, status='active').first()
    if not connection:
        flash('No active partner connection found.', 'error')
        return redirect(url_for('partner_dashboard'))
    patient = User.query.get(connection.patient_id)
    permissions = get_partner_permissions(connection)
    if not permissions.messages_enabled:
        flash('Partner messaging is currently disabled.', 'info')
        return redirect(url_for('partner_dashboard'))

    if request.method == 'POST':
        message_text = request.form.get('message', '').strip()
        if message_text:
            msg = PartnerMessage(connection_id=connection.id, sender_id=user.id, receiver_id=patient.id, message=message_text)
            db.session.add(msg)
            db.session.commit()
            flash('Message sent successfully.', 'success')
            return redirect(url_for('partner_messages'))

    messages = PartnerMessage.query.filter_by(connection_id=connection.id).order_by(PartnerMessage.created_at.desc()).all()
    return render_template('partner_messages.html', user=user, patient=patient, messages=messages, connection=connection, permissions=permissions)


@app.route('/partner/notifications')
@login_required
def partner_notifications():
    user = User.query.get(session['user_id'])
    if user.user_type != 'partner':
        flash('Access denied.', 'error')
        return redirect(url_for('index'))
    notifications = PartnerNotification.query.filter_by(user_id=user.id).order_by(PartnerNotification.created_at.desc()).all()
    return render_template('partner_notifications.html', user=user, notifications=notifications)


@app.route('/partner/sharing', methods=['GET', 'POST'])
@login_required
def partner_sharing():
    user = User.query.get(session['user_id'])
    if user is None or user.user_type != 'patient':
        flash('Only the patient can update privacy settings.', 'error')
        return redirect(url_for('patient_dashboard'))

    connection = PartnerConnection.query.filter_by(patient_id=user.id, status='active').first()
    if not connection:
        flash('You need an active partner connection to manage sharing settings.', 'info')
        return redirect(url_for('patient_dashboard'))

    permissions = get_partner_permissions(connection)
    if request.method == 'POST':
        permissions.treatment_shared = 'treatment_shared' in request.form
        permissions.appointments_shared = 'appointments_shared' in request.form
        permissions.medications_shared = 'medications_shared' in request.form
        permissions.wellness_shared = 'wellness_shared' in request.form
        permissions.mood_shared = 'mood_shared' in request.form
        permissions.nutrition_shared = 'nutrition_shared' in request.form
        permissions.documents_shared = 'documents_shared' in request.form
        permissions.doctor_notes_shared = 'doctor_notes_shared' in request.form
        permissions.messages_enabled = 'messages_enabled' in request.form
        permissions.tasks_enabled = 'tasks_enabled' in request.form
        permissions.updated_at = datetime.utcnow()
        db.session.commit()
        flash('Sharing settings updated successfully.', 'success')
        log_action(user.id, 'PARTNER_SHARING_UPDATED', 'Patient updated privacy settings')
        return redirect(url_for('partner_sharing'))

    return render_template('partner_sharing.html', user=user, connection=connection, permissions=permissions)


@app.route('/partner/disconnect', methods=['POST'])
@login_required
def partner_disconnect():
    user = User.query.get(session['user_id'])
    if user is None or user.user_type != 'patient':
        flash('Access denied.', 'error')
        return redirect(url_for('index'))
    connection = PartnerConnection.query.filter_by(patient_id=user.id, status='active').first()
    if not connection:
        flash('No active partner connection found.', 'info')
        return redirect(url_for('patient_dashboard'))
    connection.status = 'disconnected'
    connection.disconnected_at = datetime.utcnow()
    db.session.commit()
    flash('Your partner connection has been disconnected.', 'info')
    log_action(user.id, 'PARTNER_DISCONNECTED', f'Patient disconnected partner id {connection.partner_id}')
    return redirect(url_for('patient_dashboard'))


@app.route('/partner/checkin', methods=['GET', 'POST'])
@login_required
def partner_checkin():
    user = User.query.get(session['user_id'])
    if user.user_type != 'partner':
        flash('Access denied.', 'error')
        return redirect(url_for('index'))
    connection = PartnerConnection.query.filter_by(partner_id=user.id, status='active').first()
    if not connection:
        flash('No active partner connection found.', 'error')
        return redirect(url_for('partner_dashboard'))
    if request.method == 'POST':
        checkin_type = request.form.get('checkin_type', 'Emotional Support')
        note = request.form.get('note', '').strip()
        checkin = PartnerCheckIn(connection_id=connection.id, checkin_type=checkin_type, note=note)
        db.session.add(checkin)
        db.session.commit()
        flash('Support check-in saved.', 'success')
        return redirect(url_for('partner_checkin'))

    checkins = PartnerCheckIn.query.filter_by(connection_id=connection.id).order_by(PartnerCheckIn.created_at.desc()).limit(10).all()
    return render_template('partner_checkin.html', user=user, connection=connection, checkins=checkins)


@app.route('/partner/profile')
@login_required
def partner_profile():
    user = User.query.get(session['user_id'])
    if user.user_type != 'partner':
        flash('Access denied.', 'error')
        return redirect(url_for('index'))
    connection = PartnerConnection.query.filter_by(partner_id=user.id, status='active').first()
    patient = User.query.get(connection.patient_id) if connection else None
    return render_template('partner_profile.html', user=user, patient=patient, connection=connection)


@app.route('/partner/notifications/read/<int:notification_id>', methods=['POST'])
@login_required
def mark_partner_notification_read(notification_id):
    user = User.query.get(session['user_id'])
    note = PartnerNotification.query.get_or_404(notification_id)
    if note.user_id != user.id:
        flash('Access denied.', 'error')
        return redirect(url_for('partner_notifications'))
    note.is_read = True
    db.session.commit()
    return redirect(url_for('partner_notifications'))


# --- Community Forum Routes ---
@app.route('/forum')
@login_required
def forum():
    posts = ForumPost.query.order_by(ForumPost.created_at.desc()).all()
    return render_template('forum.html', posts=posts)

@app.route('/forum/new', methods=['GET', 'POST'])
@login_required
def new_post():
    if request.method == 'POST':
        title = request.form.get('title')
        content = request.form.get('content')
        if title and content:
            post = ForumPost(title=title, content=content, user_id=session['user_id'])
            db.session.add(post)
            db.session.commit()
            flash('Post created successfully!', 'success')
            return redirect(url_for('forum'))
    return render_template('new_post.html')

@app.route('/forum/post/<int:post_id>')
@login_required
def view_post(post_id):
    post = ForumPost.query.get_or_404(post_id)
    return render_template('view_post.html', post=post)

# --- Static Pages for Business & Research Framing ---
@app.route('/privacy_policy')
def privacy_policy():
    return render_template('privacy_policy.html')

@app.route('/business_model')
def business_model():
    return render_template('business_model.html')

@app.route('/roadmap')
def roadmap():
    return render_template('roadmap.html')

@app.route('/data_security')
def data_security():
    """Renders the Data Security Architecture page for research/compliance visibility."""
    return render_template('data_security.html')

# --- Admin Model Retraining Route ---
@app.route('/admin/retrain_model', methods=['POST'])
@admin_required
def retrain_model():
    try:
        from train_models import train_ivf_model
        success = train_ivf_model()
        if success:
            # Reload the model in the current app context
            global ivf_model
            ivf_model = joblib.load("models/ivf_success_model.pkl")
            flash('AI Model retrained successfully using latest patient data!', 'success')
        else:
            flash('Model training failed.', 'error')
    except Exception as e:
        app.logger.error(f"Training error: {e}")
        flash(f'Error during training: {str(e)}', 'error')
    
    return redirect(url_for('admin_dashboard'))

if __name__ == "__main__":
    # Runs on all interfaces, enabling access from 127.0.0.1:5000 and local IP
    app.run(host="0.0.0.0", port=5000, debug=True)
